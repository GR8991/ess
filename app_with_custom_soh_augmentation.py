# BESS LCOS Analysis - Streamlit App with User-Configurable SOH & Augmentation
# Version 2.0 - Professional Financial Analysis & Reporting
# FEATURES: User inputs SOH curve, Augmentation years/amounts, advanced controls

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime, timedelta
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page Configuration
st.set_page_config(
    page_title="BESS LCOS Analysis Dashboard",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main {
        background-color: #f5f5f5;
    }
    .stMetric {
        background-color: white;
        padding: 10px;
        border-radius: 5px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    h1, h2, h3 {
        color: #003366;
    }
    .section-header {
        background: linear-gradient(90deg, #003366, #0055aa);
        color: white;
        padding: 15px;
        border-radius: 5px;
        margin: 20px 0 10px 0;
    }
    .degradation-help {
        background-color: #e3f2fd;
        padding: 12px;
        border-left: 4px solid #1976d2;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# DATA INITIALIZATION & INPUT MODELS
# ============================================================================

class DegradationProfile:
    """Battery State of Health degradation - USER CONFIGURABLE"""
    def __init__(self, formation_loss=0.05, annual_degradation=0.025, floor_soh=0.60):
        self.formation_loss = formation_loss
        self.annual_degradation = annual_degradation
        self.floor_soh = floor_soh
        
    def calculate_soh(self, year):
        """Calculate SOH for given year"""
        if year == 0:
            return 1.0
        elif year == 1:
            return 1.0 - self.formation_loss
        else:
            soh = (1.0 - self.formation_loss) - (self.annual_degradation * (year - 1))
            return max(soh, self.floor_soh)
    
    def get_soh_curve(self, years):
        """Get SOH values for all years"""
        return np.array([self.calculate_soh(y) for y in years])

class CostComponents:
    """CAPEX and OPEX calculations"""
    def __init__(self, battery_cost_kwh):
        self.battery_cost_kwh = battery_cost_kwh
        self.bop_premium = 0.05  # 5% BoP
        self.epc_premium = 0.08  # 8% EPC
        self.system_cost_per_mwh = battery_cost_kwh * 1000 * (1 + self.bop_premium) * (1 + self.epc_premium)
        
    def fixed_opex_base(self):
        """Annual fixed OPEX components"""
        personnel = 300000
        land_lease = 50000
        maintenance = 400000
        monitoring = 100000
        return personnel + land_lease + maintenance + monitoring
    
    def insurance(self, initial_capex):
        """Annual insurance as % of CAPEX"""
        return initial_capex * 1000000 * 0.0075
    
    def property_tax(self, initial_capex):
        """Annual property tax as % of CAPEX"""
        return initial_capex * 1000000 * 0.005
    
    def variable_opex(self, year):
        """Age-dependent variable OPEX"""
        if year <= 5:
            return 100000
        elif year <= 10:
            return 240000
        elif year <= 15:
            return 400000
        else:
            return 550000

class RevenueModel:
    """Revenue calculations"""
    def __init__(self, power_mw, energy_margin):
        self.power_mw = power_mw
        self.energy_margin = energy_margin
        
    def capacity_payment_year(self, year):
        """Tiered capacity payments"""
        if year <= 5:
            return 50000
        elif year <= 10:
            return 50000
        elif year <= 15:
            return 45000
        else:
            return 40000
    
    def calculate_capacity_revenue(self, year):
        """Annual capacity revenue"""
        return (self.power_mw * self.capacity_payment_year(year)) / 1000000
    
    def calculate_energy_revenue(self, available_mwh, cycles_per_year, margin):
        """Annual energy arbitrage revenue"""
        return (available_mwh * cycles_per_year * margin) / 1000000

# ============================================================================
# STREAMLIT APP STRUCTURE
# ============================================================================

def main():
    st.title("⚡ BESS LCOS Analysis Dashboard")
    st.markdown("Professional Battery Energy Storage System Financial Modeling with Custom SOH & Augmentation")
    
    # Sidebar Configuration
    with st.sidebar:
        st.header("🔧 Configuration")
        analysis_type = st.radio(
            "Select Analysis Type:",
            ["Input & Analysis", "SOH Curves", "Compare Scenarios", "Generate Reports"]
        )
    
    if analysis_type == "Input & Analysis":
        show_input_analysis()
    elif analysis_type == "SOH Curves":
        show_soh_configuration()
    elif analysis_type == "Compare Scenarios":
        show_scenario_comparison()
    else:
        show_reports()

def show_soh_configuration():
    """Configure SOH curves"""
    st.markdown('<div class="section-header">📊 Battery Degradation (SOH) Configuration</div>', unsafe_allow_html=True)
    
    st.markdown("""
    <div class="degradation-help">
    <b>📌 How to Configure SOH Curve:</b><br>
    • <b>Formation Loss (Year 1):</b> Initial loss during first year (e.g., 5% = 95% SOH at end of Year 1)<br>
    • <b>Annual Degradation:</b> Loss per year after Year 1 (e.g., 2.5% per year)<br>
    • <b>Minimum SOH Floor:</b> Battery never degrades below this level (e.g., 60% minimum usable capacity)<br>
    <br>
    <b>Example:</b> LiFePO4 batteries typically have 5% formation loss, 2-3% annual degradation, 60% floor
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Degradation Parameters")
        formation_loss = st.slider(
            "Formation Loss - Year 1 (%)",
            min_value=0, max_value=10, value=5, step=1
        ) / 100
        
        annual_deg = st.slider(
            "Annual Degradation Rate (%)",
            min_value=0.5, max_value=5.0, value=2.5, step=0.5
        ) / 100
        
        floor_soh = st.slider(
            "Minimum SOH Floor (%)",
            min_value=40, max_value=80, value=60, step=5
        ) / 100
    
    with col2:
        st.subheader("SOH Curve Preview")
        
        # Create degradation profile
        degradation = DegradationProfile(
            formation_loss=formation_loss,
            annual_degradation=annual_deg,
            floor_soh=floor_soh
        )
        
        # Generate curve
        years = np.arange(0, 21)
        soh_values = degradation.get_soh_curve(years) * 100
        
        # Plot
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=years, y=soh_values,
            mode='lines+markers',
            name='SOH %',
            line=dict(color='#1976d2', width=3),
            marker=dict(size=8)
        ))
        fig.add_hline(y=floor_soh*100, line_dash="dash", line_color="red", 
                     annotation_text=f"Floor: {floor_soh*100:.0f}%")
        fig.update_layout(
            title="Battery State of Health Over Time",
            xaxis_title="Year",
            yaxis_title="SOH (%)",
            height=400,
            hovermode='x unified'
        )
        st.plotly_chart(fig, use_container_width=True)
    
    # SOH Table
    st.subheader("Detailed SOH Values")
    soh_df = pd.DataFrame({
        "Year": years,
        "SOH (%)": soh_values.round(2)
    })
    st.dataframe(soh_df, use_container_width=True)
    
    st.info("✅ Use these SOH parameters in your analysis below")

def show_input_analysis():
    """Main analysis tab with augmentation configuration"""
    st.markdown('<div class="section-header">📊 Project Configuration & Financial Analysis</div>', unsafe_allow_html=True)
    
    # Create tabs for different configuration sections
    tab_project, tab_financial, tab_augmentation, tab_soh = st.tabs([
        "Project", "Financial", "Augmentation", "Degradation"
    ])
    
    with tab_project:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Project Specifications")
            power_mw = st.number_input("Power Capacity (MW)", value=100, min_value=10, max_value=500)
            energy_mwh = st.number_input("Energy Capacity (MWh)", value=400, min_value=50, max_value=2000)
        with col2:
            st.subheader("Operational Parameters")
            project_life = st.number_input("Project Life (years)", value=20, min_value=10, max_value=40)
            cycles_per_year = st.number_input("Cycles per Year", value=365, min_value=100, max_value=1000)
    
    with tab_financial:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Battery Economics")
            battery_cost = st.number_input("Battery Cost ($/kWh)", value=241, min_value=100, max_value=500)
            cost_decline = st.slider("Annual Cost Decline Rate (%)", min_value=-8, max_value=2, value=-4, step=1) / 100
        with col2:
            st.subheader("Market & Financial")
            discount_rate = st.slider("Discount Rate (%)", min_value=3, max_value=15, value=7, step=1) / 100
            inflation_rate = st.slider("Inflation Rate (%)", min_value=0, max_value=5, value=2, step=1) / 100
            energy_margin = st.number_input("Energy Margin ($/MWh)", value=50, min_value=10, max_value=200)
    
    with tab_soh:
        st.subheader("Battery Degradation (SOH) Configuration")
        col1, col2 = st.columns(2)
        
        with col1:
            formation_loss = st.slider("Formation Loss - Year 1 (%)", 
                                      min_value=0, max_value=10, value=5, step=1) / 100
            annual_deg = st.slider("Annual Degradation Rate (%)", 
                                  min_value=0.5, max_value=5.0, value=2.5, step=0.5) / 100
        
        with col2:
            floor_soh = st.slider("Minimum SOH Floor (%)", 
                                 min_value=40, max_value=80, value=60, step=5) / 100
            
            # Preview
            temp_deg = DegradationProfile(formation_loss, annual_deg, floor_soh)
            years_preview = np.arange(0, 21)
            soh_preview = temp_deg.get_soh_curve(years_preview) * 100
            
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=years_preview, y=soh_preview, mode='lines+markers', 
                                    line=dict(color='#1976d2'), marker=dict(size=6)))
            fig.update_layout(title="SOH Preview", height=300, margin=dict(l=0, r=0, t=30, b=0))
            st.plotly_chart(fig, use_container_width=True)
    
    with tab_augmentation:
        st.subheader("Augmentation Strategy")
        st.markdown("""
        <div class="degradation-help">
        <b>📌 How to Configure Augmentation:</b><br>
        • Choose between <b>Overbuild</b> (one-time build) or <b>Staged Augmentation</b> (phased approach)<br>
        • For Staged: Specify each augmentation event with Year and Energy amount (MWh)<br>
        • Example: Year 7 → +60 MWh, Year 14 → +50 MWh
        </div>
        """, unsafe_allow_html=True)
        
        approach = st.radio("Select Approach:", ["Overbuild", "Staged Augmentation"])
        
        if approach == "Staged Augmentation":
            st.info("🔧 Configure augmentation events below")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.markdown("**Augmentation 1**")
                aug1_year = st.number_input("Year", min_value=1, max_value=project_life-2, value=7, key="aug1_year")
                aug1_capacity = st.number_input("Energy (MWh)", min_value=10, max_value=300, value=60, key="aug1_mwh")
            
            with col2:
                st.markdown("**Augmentation 2**")
                aug2_year = st.number_input("Year", min_value=aug1_year+1, max_value=project_life-1, value=14, key="aug2_year")
                aug2_capacity = st.number_input("Energy (MWh)", min_value=10, max_value=300, value=50, key="aug2_mwh")
            
            with col3:
                st.markdown("**Augmentation 3 (Optional)**")
                aug3_enable = st.checkbox("Enable 3rd augmentation?")
                if aug3_enable:
                    aug3_year = st.number_input("Year", min_value=aug2_year+1, max_value=project_life, value=18, key="aug3_year")
                    aug3_capacity = st.number_input("Energy (MWh)", min_value=10, max_value=300, value=40, key="aug3_mwh")
                else:
                    aug3_year = None
                    aug3_capacity = 0
            
            # Show augmentation timeline
            st.markdown("**Augmentation Timeline**")
            aug_events = [
                f"Year {aug1_year}: +{aug1_capacity} MWh",
                f"Year {aug2_year}: +{aug2_capacity} MWh"
            ]
            if aug3_enable and aug3_year:
                aug_events.append(f"Year {aug3_year}: +{aug3_capacity} MWh")
            
            st.success(" → ".join(aug_events))
        else:
            aug1_year = aug1_capacity = aug2_year = aug2_capacity = None
            aug3_year = aug3_capacity = None
            st.info("📌 Overbuild: Entire capacity built in Year 0")
    
    # Run Analysis
    st.markdown('<div class="section-header">🎯 Run Analysis</div>', unsafe_allow_html=True)
    
    if st.button("🚀 Calculate Financial Analysis", key="run_analysis"):
        with st.spinner("Calculating..."):
            degradation = DegradationProfile(formation_loss, annual_deg, floor_soh)
            
            if approach == "Overbuild":
                result = calculate_overbuild_cashflow(
                    power_mw, energy_mwh, project_life, cycles_per_year,
                    battery_cost, cost_decline, inflation_rate, discount_rate,
                    energy_margin, degradation
                )
                show_analysis_results("Overbuild Approach", result, "overbuild")
            else:
                result = calculate_staged_cashflow(
                    power_mw, energy_mwh, project_life, cycles_per_year,
                    battery_cost, cost_decline, inflation_rate, discount_rate,
                    energy_margin, degradation,
                    augmentations=[
                        (aug1_year, aug1_capacity),
                        (aug2_year, aug2_capacity),
                        (aug3_year, aug3_capacity) if aug3_enable else None
                    ]
                )
                show_analysis_results("Staged Augmentation Approach", result, "staged")

def show_scenario_comparison():
    """Scenario comparison tab"""
    st.markdown('<div class="section-header">📈 Scenario Analysis</div>', unsafe_allow_html=True)
    
    st.subheader("Battery Cost Decline Scenarios")
    
    scenarios = {
        "Aggressive Decline (6%/year)": -0.06,
        "Base Case (4%/year)": -0.04,
        "Moderate Decline (2%/year)": -0.02,
        "Flat / No Decline (0%/year)": 0.0,
    }
    
    scenario_results = {}
    degradation = DegradationProfile(0.05, 0.025, 0.60)
    
    for scenario_name, decline_rate in scenarios.items():
        overbuild_data = calculate_overbuild_cashflow(
            100, 400, 20, 365, 241, decline_rate, 0.02, 0.07, 50, degradation
        )
        staged_data = calculate_staged_cashflow(
            100, 400, 20, 365, 241, decline_rate, 0.02, 0.07, 50, degradation,
            augmentations=[(7, 60), (14, 50), None]
        )
        scenario_results[scenario_name] = {
            "overbuild_npv": overbuild_data["summary"]["npv"],
            "staged_npv": staged_data["summary"]["npv"],
            "advantage": staged_data["summary"]["npv"] - overbuild_data["summary"]["npv"]
        }
    
    # Create comparison dataframe
    scenario_df = pd.DataFrame([
        {
            "Scenario": name,
            "Overbuild NPV ($M)": v["overbuild_npv"],
            "Staged NPV ($M)": v["staged_npv"],
            "Staged Advantage ($M)": v["advantage"]
        }
        for name, v in scenario_results.items()
    ])
    
    st.dataframe(scenario_df, use_container_width=True)
    
    # Visualization
    fig = go.Figure()
    fig.add_trace(go.Bar(
        name="Overbuild NPV",
        x=list(scenario_results.keys()),
        y=[v["overbuild_npv"] for v in scenario_results.values()],
        marker_color="lightcoral"
    ))
    fig.add_trace(go.Bar(
        name="Staged NPV",
        x=list(scenario_results.keys()),
        y=[v["staged_npv"] for v in scenario_results.values()],
        marker_color="lightgreen"
    ))
    fig.update_layout(
        title="NPV Comparison Across Battery Cost Scenarios",
        xaxis_title="Scenario",
        yaxis_title="NPV ($M)",
        barmode="group",
        height=500
    )
    st.plotly_chart(fig, use_container_width=True)

def show_reports():
    """Reports generation tab"""
    st.markdown('<div class="section-header">📋 Generate Professional Reports</div>', unsafe_allow_html=True)
    
    st.subheader("Report Options")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        generate_csv = st.checkbox("Generate CSV Report", value=True)
    with col2:
        generate_docx = st.checkbox("Generate DOCX Report", value=True)
    with col3:
        generate_summary = st.checkbox("Generate Summary Table", value=True)
    
    if st.button("🚀 Generate All Reports"):
        degradation = DegradationProfile(0.05, 0.025, 0.60)
        overbuild_data = calculate_overbuild_cashflow(
            100, 400, 20, 365, 241, -0.04, 0.02, 0.07, 50, degradation
        )
        staged_data = calculate_staged_cashflow(
            100, 400, 20, 365, 241, -0.04, 0.02, 0.07, 50, degradation,
            augmentations=[(7, 60), (14, 50), None]
        )
        
        st.success("✅ Generating reports...")
        
        if generate_csv:
            csv_file = generate_csv_reports(overbuild_data, staged_data)
            st.download_button(
                label="📥 Download CSV Report",
                data=csv_file,
                file_name=f"BESS_LCOS_Analysis_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
        
        if generate_docx:
            docx_file = generate_docx_report(overbuild_data, staged_data)
            st.download_button(
                label="📥 Download DOCX Report",
                data=docx_file,
                file_name=f"BESS_LCOS_Analysis_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        if generate_summary:
            summary_file = generate_summary_csv(overbuild_data, staged_data)
            st.download_button(
                label="📥 Download Summary Report",
                data=summary_file,
                file_name=f"BESS_LCOS_Summary_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
        
        st.info("💡 Reports generated successfully! Use the download buttons above.")

# ============================================================================
# CALCULATION FUNCTIONS
# ============================================================================

def calculate_overbuild_cashflow(power_mw, energy_mwh, project_life, cycles,
                                 battery_cost, decline_rate, inflation, discount_rate, energy_margin,
                                 degradation):
    """Calculate overbuild scenario"""
    
    revenue_model = RevenueModel(power_mw, energy_margin)
    overbuilt_capacity = energy_mwh / degradation.floor_soh
    
    cost_component = CostComponents(battery_cost)
    initial_capex = (overbuilt_capacity * cost_component.system_cost_per_mwh) / 1000000
    
    years = np.arange(0, project_life + 1)
    cash_flows = []
    
    for year in years:
        soh = degradation.calculate_soh(year)
        available_capacity = overbuilt_capacity * soh
        
        if year == 0:
            cf = {
                "year": year,
                "soh": soh * 100,
                "capacity": available_capacity,
                "capacity_revenue": 0,
                "energy_revenue": 0,
                "total_revenue": 0,
                "fixed_opex": 0,
                "variable_opex": 0,
                "total_opex": 0,
                "capex": -initial_capex,
                "net_cf": -initial_capex
            }
        else:
            capacity_rev = revenue_model.calculate_capacity_revenue(year)
            energy_rev = revenue_model.calculate_energy_revenue(available_capacity, cycles, energy_margin)
            total_rev = capacity_rev + energy_rev
            
            cost_comp = CostComponents(battery_cost * ((1 + decline_rate) ** (year - 1)))
            fixed_opex = (cost_comp.fixed_opex_base() + cost_comp.insurance(initial_capex) + 
                         cost_comp.property_tax(initial_capex)) * (1 + inflation) ** (year - 1) / 1000000
            variable_opex = cost_comp.variable_opex(year) * (1 + inflation) ** (year - 1) / 1000000
            total_opex = fixed_opex + variable_opex
            
            cf = {
                "year": year,
                "soh": soh * 100,
                "capacity": available_capacity,
                "capacity_revenue": capacity_rev,
                "energy_revenue": energy_rev,
                "total_revenue": total_rev,
                "fixed_opex": fixed_opex,
                "variable_opex": variable_opex,
                "total_opex": total_opex,
                "capex": 0,
                "net_cf": total_rev - total_opex
            }
        
        cash_flows.append(cf)
    
    df = pd.DataFrame(cash_flows)
    df["cumulative_cf"] = df["net_cf"].cumsum()
    df["discount_factor"] = 1 / (1 + discount_rate) ** df["year"]
    df["pv_cf"] = df["net_cf"] * df["discount_factor"]
    df["cumulative_pv"] = df["pv_cf"].cumsum()
    
    total_capex = abs(df[df["year"] == 0]["net_cf"].values[0])
    total_revenue = df[df["year"] > 0]["total_revenue"].sum()
    total_opex = df[df["year"] > 0]["total_opex"].sum()
    npv = df["pv_cf"].sum()
    total_energy = df[df["year"] > 0]["capacity"].sum() * cycles
    lcos = ((total_capex + df[df["year"] > 0]["total_opex"].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    summary = {
        "initial_capex": total_capex,
        "total_revenue_nominal": total_revenue,
        "total_opex_nominal": total_opex,
        "npv": npv,
        "total_energy": total_energy,
        "lcos": lcos,
        "avg_capacity": df[df["year"] > 0]["capacity"].mean(),
        "payback_nominal": next((y for y in df["year"] if df[df["year"] == y]["cumulative_cf"].values[0] > 0), None)
    }
    
    return {"cashflow": df, "summary": summary, "overbuilt_capacity": overbuilt_capacity}

def calculate_staged_cashflow(power_mw, energy_mwh, project_life, cycles,
                             battery_cost, decline_rate, inflation, discount_rate, energy_margin,
                             degradation, augmentations=None):
    """Calculate staged augmentation scenario with custom augmentations"""
    
    revenue_model = RevenueModel(power_mw, energy_margin)
    
    # Default augmentations if not provided
    if augmentations is None:
        augmentations = [(7, 60), (14, 50), None]
    
    # Filter out None entries
    augmentations = [a for a in augmentations if a is not None]
    
    base_capacity = energy_mwh * 1.075
    
    cost_component = CostComponents(battery_cost)
    initial_capex = (base_capacity * cost_component.system_cost_per_mwh) / 1000000
    
    # Calculate augmentation costs
    aug_costs = {}
    for aug_year, aug_mwh in augmentations:
        aug_cost_kwh = battery_cost * ((1 + decline_rate) ** (aug_year - 1))
        aug_cost_comp = CostComponents(aug_cost_kwh)
        integration_premium = 1.12 if aug_year < 10 else 1.15
        aug_costs[aug_year] = (aug_mwh * aug_cost_comp.system_cost_per_mwh * integration_premium) / 1000000
    
    years = np.arange(0, project_life + 1)
    cash_flows = []
    
    for year in years:
        soh = degradation.calculate_soh(year)
        base_avail = base_capacity * soh
        
        # Calculate augmented capacity
        aug_cap = 0
        for aug_year, aug_mwh in augmentations:
            if year >= aug_year:
                aug_soh = degradation.calculate_soh(year - aug_year)
                aug_cap += aug_mwh * aug_soh
        
        total_capacity = base_avail + aug_cap
        
        if year == 0:
            cf = {
                "year": year,
                "soh": soh * 100,
                "capacity": base_capacity,
                "total_capacity": base_capacity,
                "capacity_revenue": 0,
                "energy_revenue": 0,
                "total_revenue": 0,
                "fixed_opex": 0,
                "variable_opex": 0,
                "total_opex": 0,
                "capex": -initial_capex,
                "net_cf": -initial_capex
            }
        else:
            capacity_rev = revenue_model.calculate_capacity_revenue(year)
            energy_rev = revenue_model.calculate_energy_revenue(total_capacity, cycles, energy_margin)
            total_rev = capacity_rev + energy_rev
            
            cost_comp = CostComponents(battery_cost * ((1 + decline_rate) ** (year - 1)))
            fixed_opex = (cost_comp.fixed_opex_base() + cost_comp.insurance(initial_capex) +
                         cost_comp.property_tax(initial_capex)) * (1 + inflation) ** (year - 1) / 1000000
            variable_opex = cost_comp.variable_opex(year) * (1 + inflation) ** (year - 1) / 1000000
            total_opex = fixed_opex + variable_opex
            
            # Augmentation CAPEX
            aug_capex = 0
            for aug_year, aug_mwh in augmentations:
                if year == aug_year:
                    aug_capex += -aug_costs.get(aug_year, 0)
            
            cf = {
                "year": year,
                "soh": soh * 100,
                "capacity": total_capacity,
                "total_capacity": total_capacity,
                "capacity_revenue": capacity_rev,
                "energy_revenue": energy_rev,
                "total_revenue": total_rev,
                "fixed_opex": fixed_opex,
                "variable_opex": variable_opex,
                "total_opex": total_opex,
                "capex": aug_capex,
                "net_cf": total_rev - total_opex + aug_capex
            }
        
        cash_flows.append(cf)
    
    df = pd.DataFrame(cash_flows)
    df["cumulative_cf"] = df["net_cf"].cumsum()
    df["discount_factor"] = 1 / (1 + discount_rate) ** df["year"]
    df["pv_cf"] = df["net_cf"] * df["discount_factor"]
    df["cumulative_pv"] = df["pv_cf"].cumsum()
    
    total_capex = abs(df[df["year"] == 0]["net_cf"].values[0]) + sum(aug_costs.values())
    total_revenue = df[df["year"] > 0]["total_revenue"].sum()
    total_opex = df[df["year"] > 0]["total_opex"].sum()
    npv = df["pv_cf"].sum()
    total_energy = df[df["year"] > 0]["capacity"].sum() * cycles
    lcos = ((initial_capex + sum([c / (1 + discount_rate) ** y for y, c in aug_costs.items()]) + 
             df[df["year"] > 0]["total_opex"].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    summary = {
        "initial_capex": initial_capex,
        "total_revenue_nominal": total_revenue,
        "total_opex_nominal": total_opex,
        "npv": npv,
        "total_energy": total_energy,
        "lcos": lcos,
        "avg_capacity": df[df["year"] > 0]["capacity"].mean(),
        "payback_nominal": next((y for y in df["year"] if df[df["year"] == y]["cumulative_cf"].values[0] > 0), None)
    }
    
    return {"cashflow": df, "summary": summary, "base_capacity": base_capacity, "aug_costs": aug_costs}

# ============================================================================
# VISUALIZATION & REPORTING FUNCTIONS (Simplified - same as before)
# ============================================================================

def show_analysis_results(title, data, approach_type):
    """Display analysis results with visualizations"""
    
    st.markdown(f'<div class="section-header">{title}</div>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Initial CAPEX ($M)", f"${data['summary']['initial_capex']:.2f}")
    with col2:
        st.metric("Project NPV ($M)", f"${data['summary']['npv']:.2f}")
    with col3:
        st.metric("LCOS ($/MWh)", f"${data['summary']['lcos']:.2f}")
    with col4:
        st.metric("Avg Capacity (MWh)", f"{data['summary']['avg_capacity']:.1f}")
    
    tab1, tab2, tab3, tab4 = st.tabs(["Cash Flow", "Charts", "Capacity", "Detailed Table"])
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            df = data["cashflow"]
            fig.add_trace(go.Scatter(x=df["year"], y=df["cumulative_cf"], name="Cumulative CF (Nominal)", 
                                    mode="lines+markers", line=dict(color="blue", width=3)))
            fig.add_hline(y=0, line_dash="dash", line_color="red")
            fig.update_layout(title="Cumulative Cash Flow", xaxis_title="Year", yaxis_title="$ Millions", height=400)
            st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df["year"], y=df["cumulative_pv"], name="Cumulative PV (NPV)", 
                                    mode="lines+markers", line=dict(color="darkblue", width=3)))
            fig.add_hline(y=0, line_dash="dash", line_color="red")
            fig.update_layout(title="Cumulative Present Value", xaxis_title="Year", yaxis_title="$ Millions", height=400)
            st.plotly_chart(fig, use_container_width=True)
    
    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            df_ops = df[df["year"] > 0]
            fig.add_trace(go.Bar(x=df_ops["year"], y=df_ops["total_revenue"], name="Revenue", marker_color="green"))
            fig.add_trace(go.Bar(x=df_ops["year"], y=df_ops["total_opex"], name="OPEX", marker_color="red"))
            fig.update_layout(title="Annual Revenue vs OPEX", xaxis_title="Year", yaxis_title="$ Millions", barmode="group", height=400)
            st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df["year"], y=df["soh"], name="State of Health (%)", 
                                    mode="lines+markers", line=dict(color="purple", width=3)))
            fig.update_layout(title="Battery Degradation Over Time", xaxis_title="Year", yaxis_title="SOH (%)", height=400)
            st.plotly_chart(fig, use_container_width=True)
    
    with tab3:
        fig = go.Figure()
        fig.add_trace(go.Scatter(x=df["year"], y=df["capacity"], name="Available Capacity", 
                                mode="lines+markers", line=dict(color="orange", width=3), fill="tozeroy"))
        fig.update_layout(title="Available Capacity Over Time", xaxis_title="Year", yaxis_title="MWh", height=500)
        st.plotly_chart(fig, use_container_width=True)
    
    with tab4:
        st.dataframe(df.round(2), use_container_width=True)
        csv = df.to_csv(index=False)
        st.download_button(label="📥 Download Detailed Cash Flow CSV", data=csv,
                          file_name=f"cashflow_{approach_type}_{datetime.now().strftime('%Y%m%d')}.csv", mime="text/csv")

def generate_csv_reports(overbuild_data, staged_data):
    """Generate comprehensive CSV report"""
    output = io.StringIO()
    output.write("BESS LCOS ANALYSIS - PROFESSIONAL REPORT\n")
    output.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    output.write("EXECUTIVE SUMMARY\n" + "="*80 + "\n\n")
    output.write("Metric,Overbuild,Staged,Difference\n")
    output.write(f"Initial CAPEX ($M),{overbuild_data['summary']['initial_capex']:.2f},{staged_data['summary']['initial_capex']:.2f},{staged_data['summary']['initial_capex']-overbuild_data['summary']['initial_capex']:.2f}\n")
    output.write(f"Project NPV ($M),{overbuild_data['summary']['npv']:.2f},{staged_data['summary']['npv']:.2f},{staged_data['summary']['npv']-overbuild_data['summary']['npv']:.2f}\n")
    output.write(f"LCOS ($/MWh),{overbuild_data['summary']['lcos']:.2f},{staged_data['summary']['lcos']:.2f},{staged_data['summary']['lcos']-overbuild_data['summary']['lcos']:.2f}\n")
    output.write(f"Total Revenue ($M),{overbuild_data['summary']['total_revenue_nominal']:.2f},{staged_data['summary']['total_revenue_nominal']:.2f},{staged_data['summary']['total_revenue_nominal']-overbuild_data['summary']['total_revenue_nominal']:.2f}\n")
    output.write(f"Total OPEX ($M),{overbuild_data['summary']['total_opex_nominal']:.2f},{staged_data['summary']['total_opex_nominal']:.2f},{staged_data['summary']['total_opex_nominal']-overbuild_data['summary']['total_opex_nominal']:.2f}\n\n")
    output.write("\nOVERBUILD APPROACH - YEAR BY YEAR CASH FLOW\n" + "="*80 + "\n")
    overbuild_data["cashflow"].to_csv(output, index=False)
    output.write("\n\nSTAGED AUGMENTATION APPROACH - YEAR BY YEAR CASH FLOW\n" + "="*80 + "\n")
    staged_data["cashflow"].to_csv(output, index=False)
    return output.getvalue()

def generate_summary_csv(overbuild_data, staged_data):
    """Generate summary CSV file"""
    output = io.StringIO()
    output.write("BESS LCOS Analysis - Executive Summary\n")
    output.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    summary_data = {
        "Metric": ["Initial CAPEX ($M)", "Project NPV ($M)", "LCOS ($/MWh)", "Total Revenue ($M)", 
                  "Total OPEX ($M)", "Average Capacity (MWh)", "Payback Period (years)", "Total Energy (MWh)"],
        "Overbuild": [f"{overbuild_data['summary']['initial_capex']:.2f}", f"{overbuild_data['summary']['npv']:.2f}",
                     f"{overbuild_data['summary']['lcos']:.2f}", f"{overbuild_data['summary']['total_revenue_nominal']:.2f}",
                     f"{overbuild_data['summary']['total_opex_nominal']:.2f}", f"{overbuild_data['summary']['avg_capacity']:.1f}",
                     f"{overbuild_data['summary']['payback_nominal'] if overbuild_data['summary']['payback_nominal'] else 'N/A'}",
                     f"{overbuild_data['summary']['total_energy']:.0f}"],
        "Staged": [f"{staged_data['summary']['initial_capex']:.2f}", f"{staged_data['summary']['npv']:.2f}",
                  f"{staged_data['summary']['lcos']:.2f}", f"{staged_data['summary']['total_revenue_nominal']:.2f}",
                  f"{staged_data['summary']['total_opex_nominal']:.2f}", f"{staged_data['summary']['avg_capacity']:.1f}",
                  f"{staged_data['summary']['payback_nominal'] if staged_data['summary']['payback_nominal'] else 'N/A'}",
                  f"{staged_data['summary']['total_energy']:.0f}"],
        "Advantage": [f"{staged_data['summary']['initial_capex'] - overbuild_data['summary']['initial_capex']:.2f}",
                     f"{staged_data['summary']['npv'] - overbuild_data['summary']['npv']:.2f}",
                     f"{staged_data['summary']['lcos'] - overbuild_data['summary']['lcos']:.2f}",
                     f"{staged_data['summary']['total_revenue_nominal'] - overbuild_data['summary']['total_revenue_nominal']:.2f}",
                     f"{staged_data['summary']['total_opex_nominal'] - overbuild_data['summary']['total_opex_nominal']:.2f}",
                     f"{staged_data['summary']['avg_capacity'] - overbuild_data['summary']['avg_capacity']:.1f}",
                     "Better" if staged_data['summary']['payback_nominal'] and overbuild_data['summary']['payback_nominal'] else "-",
                     f"{staged_data['summary']['total_energy'] - overbuild_data['summary']['total_energy']:.0f}"]
    }
    summary_df = pd.DataFrame(summary_data)
    summary_df.to_csv(output, index=False)
    return output.getvalue()

def generate_docx_report(overbuild_data, staged_data):
    """Generate DOCX report"""
    doc = Document()
    title = doc.add_heading('BESS LCOS Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timestamp_para = doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(f"NPV Advantage: ${staged_data['summary']['npv'] - overbuild_data['summary']['npv']:.2f}M")
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Overbuild'
    hdr_cells[2].text = 'Staged'
    hdr_cells[3].text = 'Difference'
    data_rows = [
        ['Initial CAPEX ($M)', f"${overbuild_data['summary']['initial_capex']:.2f}", 
         f"${staged_data['summary']['initial_capex']:.2f}", 
         f"${staged_data['summary']['initial_capex'] - overbuild_data['summary']['initial_capex']:.2f}"],
        ['Project NPV ($M)', f"${overbuild_data['summary']['npv']:.2f}", 
         f"${staged_data['summary']['npv']:.2f}",
         f"${staged_data['summary']['npv'] - overbuild_data['summary']['npv']:.2f}"],
        ['LCOS ($/MWh)', f"${overbuild_data['summary']['lcos']:.2f}",
         f"${staged_data['summary']['lcos']:.2f}",
         f"${staged_data['summary']['lcos'] - overbuild_data['summary']['lcos']:.2f}"],
        ['Total Revenue ($M)', f"${overbuild_data['summary']['total_revenue_nominal']:.2f}", 
         f"${staged_data['summary']['total_revenue_nominal']:.2f}", 
         f"${staged_data['summary']['total_revenue_nominal'] - overbuild_data['summary']['total_revenue_nominal']:.2f}"],
        ['Total OPEX ($M)', f"${overbuild_data['summary']['total_opex_nominal']:.2f}", 
         f"${staged_data['summary']['total_opex_nominal']:.2f}", 
         f"${staged_data['summary']['total_opex_nominal'] - overbuild_data['summary']['total_opex_nominal']:.2f}"],
        ['Avg Capacity (MWh)', f"{overbuild_data['summary']['avg_capacity']:.1f}", 
         f"{staged_data['summary']['avg_capacity']:.1f}", 
         f"{staged_data['summary']['avg_capacity'] - overbuild_data['summary']['avg_capacity']:.1f}"],
        ['Utilization Rate', f"{(overbuild_data['summary']['avg_capacity'] / 640 * 100):.1f}%", 
         f"{(staged_data['summary']['avg_capacity'] / staged_data['base_capacity'] * 100):.1f}%", 
         f"+{(staged_data['summary']['avg_capacity'] / staged_data['base_capacity'] * 100) - (overbuild_data['summary']['avg_capacity'] / 640 * 100):.1f}%"],
        ['Total Energy (MWh)', f"{overbuild_data['summary']['total_energy']:.0f}", 
         f"{staged_data['summary']['total_energy']:.0f}", 
         f"{staged_data['summary']['total_energy'] - overbuild_data['summary']['total_energy']:.0f}"],
    ]
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

if __name__ == "__main__":
    main()
