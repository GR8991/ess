'''import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches

# Define function to calculate key outputs based on sampled inputs
def economic_model(batt_cost_initial, batt_cost_decline, discount_rate, revenue_energy, revenue_capacity, degradation_rate):
    # Simplified model for demonstration purposes
    project_lifetime = 20
    initial_cost = batt_cost_initial * 100  # for 100 kWh base capacity
    # Assume cost declines linearly over lifetime
    cost_over_time = initial_cost * np.exp(-batt_cost_decline * np.arange(project_lifetime))
    # Calculate NPV of costs
    npv_cost = np.sum(cost_over_time / ((1 + discount_rate) ** np.arange(1, project_lifetime + 1)))
    # Revenue modeled as stable stream declining by degradation
    annual_revenue = (revenue_energy + revenue_capacity) * 100 * (1 - degradation_rate) ** np.arange(project_lifetime)
    npv_revenue = np.sum(annual_revenue / ((1 + discount_rate) ** np.arange(1, project_lifetime + 1)))
    # Calculate NPV of net cash flow
    npv = npv_revenue - npv_cost
    # Estimate payback year
    cumulative_cashflow = np.cumsum(annual_revenue - cost_over_time)
    payback = np.argmax(cumulative_cashflow > 0) + 1 if np.any(cumulative_cashflow > 0) else project_lifetime
    # LCOS = Total discounted costs / Total delivered energy (simplified)
    total_energy = 100 * project_lifetime * (1 - degradation_rate / 2)  # approx average degrade
    lcos = npv_cost / total_energy
    return npv, payback, lcos

# Run Monte Carlo simulation
def run_simulation(n_sims, params):
    results = []
    for _ in range(n_sims):
        sample = {key: np.random.uniform(low, high) for key, (low, high) in params.items()}
        npv, payback, lcos = economic_model(
            sample['batt_cost_initial'],
            sample['batt_cost_decline'],
            sample['discount_rate'],
            sample['revenue_energy'],
            sample['revenue_capacity'],
            sample['degradation_rate']
        )
        results.append({'NPV': npv, 'Payback': payback, 'LCOS': lcos, **sample})
    return pd.DataFrame(results)

def create_word_report(df, inputs):
    doc = Document()
    doc.add_heading('Monte Carlo Simulation Report - BESS Economic Model', level=1)
    doc.add_heading('Input Parameters', level=2)
    for key, val in inputs.items():
        doc.add_paragraph(f"{key}: {val}")
    doc.add_heading('Simulation Results Summary', level=2)
    desc = df.describe().transpose()
    doc.add_paragraph(desc.to_string())
    # Add plots
    fig, axs = plt.subplots(1, 3, figsize=(15, 4))
    df['NPV'].hist(ax=axs[0], bins=30, color='skyblue')
    axs[0].set_title('NPV Distribution')
    df['Payback'].hist(ax=axs[1], bins=30, color='salmon')
    axs[1].set_title('Payback Distribution')
    df['LCOS'].hist(ax=axs[2], bins=30, color='lightgreen')
    axs[2].set_title('LCOS Distribution')
    plt.tight_layout()
    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    plt.close(fig)
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(6))
    return doc

# Streamlit UI
st.title('BESS Economic Model Monte Carlo Simulation')

st.sidebar.header('Simulation Parameters')
batt_cost_initial = st.sidebar.slider('Initial Battery Cost ($/kWh)', 300, 500, (350, 450), step=10)
batt_cost_decline = st.sidebar.slider('Battery Cost Decline Rate (%/year)', 0.0, 0.06, (0.01, 0.04), step=0.005)
discount_rate = st.sidebar.slider('Discount Rate', 0.03, 0.09, (0.05, 0.08), step=0.005)
revenue_energy = st.sidebar.slider('Energy Revenue ($/kWh-year)', 30, 70, (40, 60), step=1)
revenue_capacity = st.sidebar.slider('Capacity Revenue ($/kWh-year)', 10, 50, (20, 40), step=1)
degradation_rate = st.sidebar.slider('System Degradation Rate (%/year)', 0.01, 0.04, (0.02, 0.03), step=0.001)
n_sims = st.sidebar.number_input('Number of Simulations', 1000, 10000, 5000, step=500)

params = {
    'batt_cost_initial': batt_cost_initial,
    'batt_cost_decline': batt_cost_decline,
    'discount_rate': discount_rate,
    'revenue_energy': revenue_energy,
    'revenue_capacity': revenue_capacity,
    'degradation_rate': degradation_rate
}

if st.button('Run Simulation'):
    with st.spinner('Running simulations...'):
        df_results = run_simulation(n_sims, params)

    st.success('Simulation complete!')
    st.subheader('Simulation Results')
    st.write(df_results.describe())

    fig, axs = plt.subplots(1, 3, figsize=(15, 4))
    df_results['NPV'].hist(ax=axs[0], bins=30, color='skyblue')
    axs[0].set_title('NPV Distribution')
    df_results['Payback'].hist(ax=axs[1], bins=30, color='salmon')
    axs[1].set_title('Payback Distribution')
    df_results['LCOS'].hist(ax=axs[2], bins=30, color='lightgreen')
    axs[2].set_title('LCOS Distribution')
    st.pyplot(fig)

    # Download Word report
    doc = create_word_report(df_results, {k: f"{v[0]:.3f} to {v[1]:.3f}" for k, v in params.items()})
    buf = BytesIO()
    doc.save(buf)
    st.download_button(
        label="Download Word Report",
        data=buf.getvalue(),
        file_name="BESS_MonteCarlo_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )-----------------
import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Define function to calculate key outputs based on sampled inputs
def economic_model(batt_cost_initial, batt_cost_decline, discount_rate, revenue_energy, revenue_capacity, degradation_rate):
    project_lifetime = 20
    initial_cost = batt_cost_initial * 100  # for 100 kWh base capacity
    cost_over_time = initial_cost * np.exp(-batt_cost_decline * np.arange(project_lifetime))
    npv_cost = np.sum(cost_over_time / ((1 + discount_rate) ** np.arange(1, project_lifetime + 1)))
    annual_revenue = (revenue_energy + revenue_capacity) * 100 * (1 - degradation_rate) ** np.arange(project_lifetime)
    npv_revenue = np.sum(annual_revenue / ((1 + discount_rate) ** np.arange(1, project_lifetime + 1)))
    npv = npv_revenue - npv_cost
    cumulative_cashflow = np.cumsum(annual_revenue - cost_over_time)
    payback = np.argmax(cumulative_cashflow > 0) + 1 if np.any(cumulative_cashflow > 0) else project_lifetime
    total_energy = 100 * project_lifetime * (1 - degradation_rate / 2)
    lcos = npv_cost / total_energy
    return npv, payback, lcos

# Run Monte Carlo simulation
def run_simulation(n_sims, params):
    results = []
    for _ in range(n_sims):
        sample = {key: np.random.uniform(low, high) for key, (low, high) in params.items()}
        npv, payback, lcos = economic_model(
            sample['batt_cost_initial'],
            sample['batt_cost_decline'],
            sample['discount_rate'],
            sample['revenue_energy'],
            sample['revenue_capacity'],
            sample['degradation_rate']
        )
        results.append({'NPV': npv, 'Payback': payback, 'LCOS': lcos, **sample})
    return pd.DataFrame(results)

def create_word_report(df, inputs, opinion):
    doc = Document()
    doc.add_heading('Monte Carlo Simulation Report - BESS Economic Model', level=1)
    doc.add_heading('Input Parameters', level=2)
    for key, val in inputs.items():
        doc.add_paragraph(f"{key}: {val}")
    doc.add_heading('Simulation Results Summary', level=2)

    # Add summary table
    desc = df.describe().transpose()
    table = doc.add_table(rows=1, cols=len(desc.columns)+1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    for i, colname in enumerate(desc.columns):
        hdr_cells[i+1].text = colname.capitalize()
    for metric, row in desc.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(metric)
        for i, val in enumerate(row):
            row_cells[i+1].text = f"{val:.2f}"

    doc.add_heading('Interpretation for Non-Economists', level=2)
    p = doc.add_paragraph(opinion)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add plots
    fig, axs = plt.subplots(1, 3, figsize=(15, 4))
    df['NPV'].hist(ax=axs[0], bins=30, color='skyblue')
    axs[0].set_title('NPV Distribution')
    df['Payback'].hist(ax=axs[1], bins=30, color='salmon')
    axs[1].set_title('Payback Distribution')
    df['LCOS'].hist(ax=axs[2], bins=30, color='lightgreen')
    axs[2].set_title('LCOS Distribution')
    plt.tight_layout()
    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    plt.close(fig)
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(6))
    
    return doc

# Generate plain language opinion based on results summary
def generate_opinion(df):
    mean_npv = df['NPV'].mean()
    mean_payback = df['Payback'].mean()
    mean_lcos = df['LCOS'].mean()
    opinion = (f"The average Net Present Value (NPV) across simulations is ${mean_npv:,.2f}, "
               f"which indicates the project is generally profitable. "
               f"The average payback period is around {mean_payback:.1f} years, "
               f"meaning on average it takes this long to recover the investment. "
               f"The Levelized Cost of Storage (LCOS) averages at ${mean_lcos:,.2f} per kWh, "
               "which is a key metric to compare against alternative storage technologies or market prices. "
               "These results incorporate uncertainties in battery costs, revenues, discount rates, and degradation, "
               "providing a realistic estimate of possible economic outcomes.")
    return opinion

# Streamlit UI
st.title('BESS Economic Model Monte Carlo Simulation with Interpretation')

st.sidebar.header('Simulation Parameters')
batt_cost_initial = st.sidebar.slider('Initial Battery Cost ($/kWh)', 300, 500, (350, 450), step=10)
batt_cost_decline = st.sidebar.slider('Battery Cost Decline Rate (%/year)', 0.0, 0.06, (0.01, 0.04), step=0.005)
discount_rate = st.sidebar.slider('Discount Rate', 0.03, 0.09, (0.05, 0.08), step=0.005)
revenue_energy = st.sidebar.slider('Energy Revenue ($/kWh-year)', 30, 70, (40, 60), step=1)
revenue_capacity = st.sidebar.slider('Capacity Revenue ($/kWh-year)', 10, 50, (20, 40), step=1)
degradation_rate = st.sidebar.slider('System Degradation Rate (%/year)', 0.01, 0.04, (0.02, 0.03), step=0.001)
n_sims = st.sidebar.number_input('Number of Simulations', 1000, 10000, 5000, step=500)

params = {
    'batt_cost_initial': batt_cost_initial,
    'batt_cost_decline': batt_cost_decline,
    'discount_rate': discount_rate,
    'revenue_energy': revenue_energy,
    'revenue_capacity': revenue_capacity,
    'degradation_rate': degradation_rate
}

if st.button('Run Simulation'):
    with st.spinner('Running simulations...'):
        df_results = run_simulation(n_sims, params)

    st.success('Simulation complete!')

    st.subheader('Simulation Results Summary Table')
    st.dataframe(df_results.describe().transpose().style.format("{:.2f}"))

    opinion_text = generate_opinion(df_results)
    st.subheader('Interpretation for Non-Economists')
    st.write(opinion_text)

    fig, axs = plt.subplots(1, 3, figsize=(15, 4))
    df_results['NPV'].hist(ax=axs[0], bins=30, color='skyblue')
    axs[0].set_title('NPV Distribution')
    df_results['Payback'].hist(ax=axs[1], bins=30, color='salmon')
    axs[1].set_title('Payback Distribution')
    df_results['LCOS'].hist(ax=axs[2], bins=30, color='lightgreen')
    axs[2].set_title('LCOS Distribution')
    st.pyplot(fig)

    doc = create_word_report(df_results, {k: f"{v[0]:.3f} to {v[1]:.3f}" for k, v in params.items()}, opinion_text)
    buf = BytesIO()
    doc.save(buf)
    st.download_button(
        label="Download Word Report",
        data=buf.getvalue(),
        file_name="BESS_MonteCarlo_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )'''
import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Linear degradation function
def linear_degradation(degradation_rate, years):
    return (1 - degradation_rate) ** years

# Example nonlinear degradation function for LFP batteries:
# This is a simplified exponential decay model as an example
def nonlinear_degradation(degradation_rate, years):
    # Nonlinear degradation: faster initial degradation slowing over time
    return np.exp(-degradation_rate * years**0.5)

# Economic model with selectable degradation function
def economic_model(batt_cost_initial, batt_cost_decline, discount_rate, revenue_energy, revenue_capacity, degradation_rate, degradation_model, project_lifetime):
    initial_cost = batt_cost_initial * 100  # for 100 kWh base capacity

    years = np.arange(project_lifetime)
    cost_over_time = initial_cost * np.exp(-batt_cost_decline * years)
    if degradation_model == 'Linear':
        soh = linear_degradation(degradation_rate, years)
    else:
        soh = nonlinear_degradation(degradation_rate, years)
    
    annual_revenue = (revenue_energy + revenue_capacity) * 100 * soh
    npv_cost = np.sum(cost_over_time / ((1 + discount_rate) ** (years + 1)))
    npv_revenue = np.sum(annual_revenue / ((1 + discount_rate) ** (years + 1)))
    npv = npv_revenue - npv_cost
    
    cumulative_cashflow = np.cumsum(annual_revenue - cost_over_time)
    payback = np.argmax(cumulative_cashflow > 0) + 1 if np.any(cumulative_cashflow > 0) else project_lifetime
    
    # Approximate total energy delivered (integral of SOH curve)
    total_energy = 100 * np.sum(soh)
    lcos = npv_cost / total_energy
    
    return npv, payback, lcos

# Run Monte Carlo simulation
def run_simulation(n_sims, params):
    results = []
    for _ in range(n_sims):
        sample = {key: np.random.uniform(low, high) for key, (low, high) in params.items()}
        npv, payback, lcos = economic_model(
            sample['batt_cost_initial'],
            sample['batt_cost_decline'],
            sample['discount_rate'],
            sample['revenue_energy'],
            sample['revenue_capacity'],
            sample['degradation_rate'],
            sample['degradation_model'],
            sample['project_lifetime']
        )
        results.append({'NPV': npv, 'Payback': payback, 'LCOS': lcos, **sample})
    return pd.DataFrame(results)

def create_word_report(df, inputs, opinion):
    doc = Document()
    doc.add_heading('Monte Carlo Simulation Report - BESS Economic Model', level=1)
    doc.add_heading('Input Parameters', level=2)
    for key, val in inputs.items():
        doc.add_paragraph(f"{key}: {val}")
    doc.add_heading('Simulation Results Summary', level=2)

    desc = df.describe().transpose()
    table = doc.add_table(rows=1, cols=len(desc.columns)+1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    for i, colname in enumerate(desc.columns):
        hdr_cells[i+1].text = colname.capitalize()
    for metric, row in desc.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(metric)
        for i, val in enumerate(row):
            row_cells[i+1].text = f"{val:.2f}"

    doc.add_heading('Interpretation for Non-Economists', level=2)
    p = doc.add_paragraph(opinion)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    fig, axs = plt.subplots(1, 3, figsize=(15, 4))
    df['NPV'].hist(ax=axs[0], bins=30, color='skyblue')
    axs[0].set_title('NPV Distribution')
    df['Payback'].hist(ax=axs[1], bins=30, color='salmon')
    axs[1].set_title('Payback Distribution')
    df['LCOS'].hist(ax=axs[2], bins=30, color='lightgreen')
    axs[2].set_title('LCOS Distribution')
    plt.tight_layout()
    
    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    plt.close(fig)
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(6))
    
    return doc

def generate_opinion(df):
    mean_npv = df['NPV'].mean()
    mean_payback = df['Payback'].mean()
    mean_lcos = df['LCOS'].mean()
    opinion = (
        f"The average Net Present Value (NPV) is ${mean_npv:,.2f}, indicating the project is generally profitable. "
        f"The average payback period is {mean_payback:.1f} years, i.e., time to recover investment costs. "
        f"The Levelized Cost of Storage (LCOS) averages ${mean_lcos:,.2f} per kWh, which is a key economic metric. "
        "Results factor in uncertainties in costs, revenues, discount rate, and degradation, offering realistic economic outcome estimates."
    )
    return opinion

# Streamlit UI
st.title('BESS Economic Model Monte Carlo Simulation with Degradation Options')

st.sidebar.header('Simulation Parameters')
batt_cost_initial = st.sidebar.slider('Initial Battery Cost ($/kWh)', 300, 500, (350, 450), step=10)
batt_cost_decline = st.sidebar.slider('Battery Cost Decline Rate (%/year)', 0.0, 0.06, (0.01, 0.04), step=0.005)
discount_rate = st.sidebar.slider('Discount Rate', 0.03, 0.09, (0.05, 0.08), step=0.005)
revenue_energy = st.sidebar.slider('Energy Revenue ($/kWh-year)', 30, 70, (40, 60), step=1)
revenue_capacity = st.sidebar.slider('Capacity Revenue ($/kWh-year)', 10, 50, (20, 40), step=1)
degradation_rate = st.sidebar.slider('Degradation Rate', 0.01, 0.04, (0.02, 0.03), step=0.001)
degradation_model = st.sidebar.radio('Degradation Model', ['Linear', 'Nonlinear'])
project_lifetime = st.sidebar.slider('Project Lifetime (years)', 5, 30, 20, step=1)

n_sims = st.sidebar.number_input('Number of Simulations', 1000, 10000, 5000, step=500)

params = {
    'batt_cost_initial': batt_cost_initial,
    'batt_cost_decline': batt_cost_decline,
    'discount_rate': discount_rate,
    'revenue_energy': revenue_energy,
    'revenue_capacity': revenue_capacity,
    'degradation_rate': degradation_rate,
    'degradation_model': degradation_model,
    'project_lifetime': project_lifetime
}

if st.button('Run Simulation'):
    with st.spinner('Running simulations...'):
        df_results = run_simulation(n_sims, params)

    st.success('Simulation complete!')

    st.subheader('Simulation Results Summary Table')
    st.dataframe(df_results.describe().transpose().style.format("{:.2f}"))

    opinion_text = generate_opinion(df_results)
    st.subheader('Interpretation for Non-Economists')
    st.write(opinion_text)

    fig, axs = plt.subplots(1, 3, figsize=(15, 4))
    df_results['NPV'].hist(ax=axs[0], bins=30, color='skyblue')
    axs[0].set_title('NPV Distribution')
    df_results['Payback'].hist(ax=axs[1], bins=30, color='salmon')
    axs[1].set_title('Payback Distribution')
    df_results['LCOS'].hist(ax=axs[2], bins=30, color='lightgreen')
    axs[2].set_title('LCOS Distribution')
    st.pyplot(fig)

    doc = create_word_report(df_results, {
        k: f"{v[0]:.3f} to {v[1]:.3f}" if isinstance(v, tuple) else str(v) for k, v in params.items()
    }, opinion_text)
    buf = BytesIO()
    doc.save(buf)
    st.download_button(
        label="Download Word Report",
        data=buf.getvalue(),
        file_name="BESS_MonteCarlo_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


