import streamlit as st
import math
from docx import Document
from io import BytesIO
import smtplib
from email.message import EmailMessage

st.set_page_config(page_title="PGR-Earthing Design Calculator", layout="centered")
st.title("⚡ PGR - Earthing Design Calculator")

# Input Section
st.sidebar.header("Input Parameters")

# Ask user for type of earth conductor
earth_conductor_type = st.sidebar.radio("Select Earth Electrode Type:", ["MS/GI Flat", "Rod"])

if earth_conductor_type == "MS/GI Flat":
    electrode_width = st.sidebar.number_input("Electrode Width (mm)", value=50.0)
    electrode_thickness = st.sidebar.number_input("Electrode Thickness (mm)", value=6.0)
    dia = (math.sqrt(4 * electrode_width * electrode_thickness / math.pi)) / 1000
else:
    rod_dia = st.sidebar.number_input("Rod Electrode Diameter (mm)", value=25.0)
    dia = rod_dia / 1000

inputs = {
    "Earth Resistance (Ω)": st.sidebar.number_input("Earth Resistance (Ω)", value=1.0),
    "Surface Type": st.sidebar.selectbox("Surface Type", ["Crushed", "Asphalt"]),
    "Layer Fault Duration (s)": st.sidebar.number_input("Layer Fault Duration (s)", value=1.0),
    "Layer Thickness (m)": st.sidebar.number_input("Layer Thickness (m)", value=0.15),
    "Df - Decrement Factor": st.sidebar.number_input("Df - Decrement Factor", value=1.3),
    "Sf - Current Division Factor": st.sidebar.number_input("Sf - Current Division Factor", value=0.5),
    "If - Ground Fault Current (kA)": st.sidebar.number_input("If - Symmetrical Ground Fault Current (kA)", value=10.0),
    "Lc - Total Earth Conductor (m)": st.sidebar.number_input("Lc - Total Earth Conductor (m)", value=300.0),
    "No. of Electrodes": st.sidebar.number_input("No. of Electrodes", min_value=1, step=1, value=30),
    "Length of Each Electrode (m)": st.sidebar.number_input("Length of Each Electrode (m)", value=3),
    "Selected Grid Spacing (m)": st.sidebar.number_input("Selected Grid Spacing (m)", value=0.1),
    "Area of Earth Mat (m²)": st.sidebar.number_input("Area of Earth Mat (m²)", value=400.0),
    "Depth of Buried Conductor (m)": st.sidebar.number_input("Depth of Buried Conductor (m)", value=0.6),
    "Perimeter of the Grid (m)": st.sidebar.number_input("Perimeter of the Grid (m)", value=80.0),
    "Max Length in X (m)": st.sidebar.number_input("Max Length in X (m)", value=20.0),
    "Max Length in Y (m)": st.sidebar.number_input("Max Length in Y (m)", value=20.0)
}

# Assign variables
c = inputs["Earth Resistance (Ω)"]
g = inputs["Surface Type"]
b = inputs["Layer Fault Duration (s)"]
d = inputs["Layer Thickness (m)"]
df = inputs["Df - Decrement Factor"]
sf = inputs["Sf - Current Division Factor"]
If = inputs["If - Ground Fault Current (kA)"]
lc = inputs["Lc - Total Earth Conductor (m)"]
ne = inputs["No. of Electrodes"]
le = inputs["Length of Each Electrode (m)"]
D = inputs["Selected Grid Spacing (m)"]
area = inputs["Area of Earth Mat (m²)"]
h = inputs["Depth of Buried Conductor (m)"]
p = inputs["Perimeter of the Grid (m)"]
lx = inputs["Max Length in X (m)"]
ly = inputs["Max Length in Y (m)"]

# Intermediate Calculations
Lr = ne * le
Lt = lc * Lr
Ls = 0.75 * lc + 0.85 * Lr
na = 2 * lc / p
nb = math.sqrt(p / (4 * math.sqrt(area)))
n = na * nb
ki = 0.644 + 0.148 * n
s = 1 / math.pi * (1 / (2 * h) + 1 / (D + h) + 1 / D * 1 - 0.5 ** (n - 2))
Ks_rounded = round(s, 4)
lm = lc + (1.55 + 1.22 * (le / math.sqrt(lx ** 2 + ly ** 2))) * Lr

km = 1 / (2 * math.pi) * (
    math.log((D ** 2) / (16 * h * dia) + ((D + 2 * h) ** 2) / (8 * D * dia) - (h / (4 * dia))) +
    1 / 1.27 * math.log(8 / (math.pi * (2 * n - 1)))
)

a = 3000 if g == "Crushed" else 10000
e = 1 - (0.09 * (1 - (c / a))) / ((2 * d) + 0.09)
f = round(e, 3)

p_70 = ((1000 + (1.5 * f * a)) * 0.157) / math.sqrt(b)
z_50 = ((1000 + (1.5 * f * a)) * 0.116) / math.sqrt(b)
x_70 = ((1000 + (6 * f * a)) * 0.157) / math.sqrt(b)
y_50 = ((1000 + (6 * f * a)) * 0.116) / math.sqrt(b)

IG = df * sf * If
Rg = (c * (1 / Lt + 1 / math.sqrt(20 * area) * (1 + 1 / (1 + h * math.sqrt(20 / area)))) *
      1.52 * (2 * math.log(p_70 * math.sqrt(2 / area) - 1) * math.sqrt(area) / p_70))
es = (c * IG * Ks_rounded * ki) / Ls * 1000
em = (c * IG * km * ki) / lm * 1000

results = {
    "Touch Voltage (50kg)": f"{round(z_50, 2)} V",
    "Touch Voltage (70kg)": f"{round(p_70, 2)} V",
    "Step Voltage (50kg)": f"{round(y_50, 2)} V",
    "Step Voltage (70kg)": f"{round(x_70, 2)} V",
    "IG (kA)": f"{IG:.2f}",
    "Rg (Ω)": f"{Rg:.2f}",
    "km (Ω)": f"{km:.2f}",
    "lm (Ω)": f"{lm:.2f}",
    "ki (Ω)": f"{ki:.2f}",
    "ls (Ω)": f"{Ls:.2f}",
    "ks (Ω)": f"{Ks_rounded:.2f}",
    "Step Potential": f"{es:.2f} V",
    "Touch Potential": f"{em:.2f} V"
}

st.subheader("📊 Results Summary")
st.table(results)

st.subheader("📧 Send Report to Email (Optional)")
receiver_email = st.text_input("Enter email to receive report (leave blank to skip):")

doc = Document()
doc.add_heading("PGR-Earthing Design Report", level=0)

doc.add_heading("Input Parameters", level=1)
table1 = doc.add_table(rows=1, cols=2)
table1.style = 'Light List Accent 1'
table1.rows[0].cells[0].text = "Parameter"
table1.rows[0].cells[1].text = "Value"

for key, val in inputs.items():
    row = table1.add_row().cells
    row[0].text = str(key)
    row[1].text = str(val)

doc.add_heading("Calculation Results", level=1)
table2 = doc.add_table(rows=1, cols=2)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = "Parameter"
table2.rows[0].cells[1].text = "Value"

for key, val in results.items():
    row = table2.add_row().cells
    row[0].text = str(key)
    row[1].text = str(val)

buffer = BytesIO()
doc.save(buffer)
buffer.seek(0)

st.download_button(
    label="📥 Download Word Report",
    data=buffer,
    file_name="earthing_report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

st.subheader("🧮 Conductor Sizing (Based on Fault Current)")

# Initialize session state
if "show_conductor_inputs" not in st.session_state:
    st.session_state.show_conductor_inputs = False

if st.button("🧾 Start Conductor Sizing"):
    st.session_state.show_conductor_inputs = True

if st.session_state.show_conductor_inputs:
    st.markdown("### 🔢 Enter Conductor Parameters")

    # Material selection and associated data
    material_data = {
        "Copper":    {"B": 234, "Qc": 3.45e-3, "delta_20": 17.241e-6},
        "Aluminium": {"B": 228, "Qc": 2.5e-3,  "delta_20": 28.264e-6},
        "Lead":      {"B": 230, "Qc": 1.45e-3, "delta_20": 214e-6},
        "Steel":     {"B": 202, "Qc": 3.8e-3,  "delta_20": 138e-6},
    }

    # Step 1: Let user select material
    selected_material = st.selectbox("Select Conductor Material", list(material_data.keys()))
    B = material_data[selected_material]["B"]
    Qc = material_data[selected_material]["Qc"]
    delta_20 = material_data[selected_material]["delta_20"]

    # Step 2: Ask for other required inputs
    Isc = st.number_input("Short-circuit current Isc (A)", value=12000.0)
    t = st.number_input("Fault duration t (s)", value=1.0)
    Tm = st.number_input("Maximum temperature Tm (°C)", value=620.0)
    Ta = st.number_input("Ambient temperature Ta (°C)", value=50.0)

    # Display auto-filled constants for transparency
    st.markdown(f"**Auto-filled Constants for {selected_material}:**")
    st.markdown(f"- B = `{B}`")
    st.markdown(f"- Qc = `{Qc}` (1/°C)")
    st.markdown(f"- Resistivity at 20°C = `{delta_20}` Ω·m")

    # Step 3: Compute conductor size
    if st.button("Calculate Size"):
        try:
            log_term = math.log((B + Tm) / (B + Ta))
            bb = Isc * math.sqrt(t) * math.sqrt(1 / ((Qc * (B + 20)) / delta_20 * log_term))
            st.success(f"✅ Required Earth Conductor Cross-Sectional Area: **{bb:.2f} mm²**")
        except Exception as e:
            st.error(f"Calculation failed: {e}")


if receiver_email:
    sender_email = "youremail@gmail.com"
    sender_password = "yourapppassword"

    try:
        msg = EmailMessage()
        msg["Subject"] = "Earthing Design Report"
        msg["From"] = sender_email
        msg["To"] = receiver_email
        msg.set_content("Please find the attached earthing design report.")

        buffer.seek(0)
        msg.add_attachment(buffer.read(), maintype="application",
                           subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                           filename="earthing_report.docx")

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(sender_email, sender_password)
            server.send_message(msg)

        st.success(f"Report successfully sent to {receiver_email} 📬")

    except Exception as e:
        st.error(f"Failed to send email: {e}")


