# BESS LCOS Analysis Dashboard - Based on Excel Sheet Structure
# Implements Sheet1, Sheet2, Sheet3, Sheet4 logic
# Version 1.0 - Production Ready

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="BESS LCOS Analysis", page_icon="⚡", layout="wide")

st.markdown("""<style>
.main { background-color: #f8f9fa; }
h1, h2, h3 { color: #003366; }
.header { background: linear-gradient(135deg, #003366 0%, #0055aa 100%); color: white; padding: 20px; border-radius: 8px; margin: 20px 0; }
</style>""", unsafe_allow_html=True)

if 'results' not in st.session_state:
    st.session_state.results = {}
if 'params' not in st.session_state:
    st.session_state.params = {}
if 'approach' not in st.session_state:
    st.session_state.approach = None

# ============================================================================
# CALCULATION ENGINE - Matches Excel Sheets 2 & 3
# ============================================================================

def calculate_overbuild(project_params, soh_curve):
    """Sheet 2: Overbuild Approach Calculation"""
    
    # Extract parameters
    project_life = project_params['project_life']
    power_mw = project_params['power_mw']
    energy_mwh = project_params['energy_mwh']
    cycles_year = project_params['cycles_year']
    battery_cost = project_params['battery_cost']
    cost_decline = project_params['cost_decline']
    inflation = project_params['inflation']
    discount_rate = project_params['discount_rate']
    energy_margin = project_params['energy_margin']
    
    # Calculate from Sheet 1
    floor_soh = min(soh_curve.values()) / 100
    overbuild_capacity = energy_mwh / floor_soh
    
    system_cost_per_mwh = battery_cost * 1000 * 1.05 * 1.08  # BoP + EPC
    initial_capex = (overbuild_capacity * system_cost_per_mwh) / 1000000
    
    # Build 20-year cash flow (Sheet 2 structure)
    cf_data = []
    
    for year in range(0, project_life + 1):
        soh_pct = soh_curve.get(year, min(soh_curve.values())) / 100
        available_capacity = overbuild_capacity * soh_pct
        
        row = {'year': year, 'soh': soh_pct * 100, 'available_capacity': available_capacity}
        
        if year == 0:
            row.update({'capacity_payment': 0, 'capacity_revenue': 0, 'energy_revenue': 0, 
                       'total_revenue': 0, 'fixed_opex': 0, 'variable_opex': 0, 'total_opex': 0,
                       'initial_capex': -initial_capex, 'aug_capex': 0, 'net_cf': -initial_capex})
        else:
            # Capacity payment tier
            if year <= 5: cp = 50000
            elif year <= 10: cp = 50000
            elif year <= 15: cp = 45000
            else: cp = 40000
            
            capacity_revenue = (power_mw * cp) / 1000000
            energy_revenue = (available_capacity * cycles_year * energy_margin) / 1000000
            total_revenue = capacity_revenue + energy_revenue
            
            fixed_opex = (850000 * (1 + inflation) ** (year - 1)) / 1000000
            
            if year <= 5: var_opex_base = 100000
            elif year <= 10: var_opex_base = 240000
            elif year <= 15: var_opex_base = 400000
            else: var_opex_base = 550000
            
            variable_opex = (var_opex_base * (1 + inflation) ** (year - 1)) / 1000000
            total_opex = fixed_opex + variable_opex
            
            net_cf = total_revenue - total_opex
            
            row.update({'capacity_payment': cp, 'capacity_revenue': capacity_revenue,
                       'energy_revenue': energy_revenue, 'total_revenue': total_revenue,
                       'fixed_opex': fixed_opex, 'variable_opex': variable_opex,
                       'total_opex': total_opex, 'initial_capex': 0, 'aug_capex': 0, 'net_cf': net_cf})
        
        cf_data.append(row)
    
    df = pd.DataFrame(cf_data)
    df['cumulative_cf'] = df['net_cf'].cumsum()
    df['discount_factor'] = 1 / (1 + discount_rate) ** df['year']
    df['pv_cf'] = df['net_cf'] * df['discount_factor']
    df['cumulative_pv'] = df['pv_cf'].cumsum()
    
    # Calculate summary metrics (Sheet 2 summary section)
    total_capex = abs(df.loc[0, 'net_cf'])
    total_revenue = df[df['year'] > 0]['total_revenue'].sum()
    total_opex = df[df['year'] > 0]['total_opex'].sum()
    npv = df['pv_cf'].sum()
    total_energy = df[df['year'] > 0]['available_capacity'].sum() * cycles_year
    lcos = ((total_capex + total_opex) * 1000000) / (total_energy * 1000)
    
    return {
        'name': 'Overbuild',
        'cashflow': df,
        'capex': total_capex,
        'npv': npv,
        'lcos': lcos,
        'total_revenue': total_revenue,
        'total_opex': total_opex,
        'avg_capacity': df[df['year'] > 0]['available_capacity'].mean(),
        'total_energy': total_energy
    }

def calculate_staged(project_params, soh_curve, augmentations):
    """Sheet 3: Staged Augmentation Calculation"""
    
    project_life = project_params['project_life']
    power_mw = project_params['power_mw']
    energy_mwh = project_params['energy_mwh']
    cycles_year = project_params['cycles_year']
    battery_cost = project_params['battery_cost']
    cost_decline = project_params['cost_decline']
    inflation = project_params['inflation']
    discount_rate = project_params['discount_rate']
    energy_margin = project_params['energy_margin']
    
    floor_soh = min(soh_curve.values()) / 100
    base_capacity = energy_mwh * 1.075  # 7.5% buffer
    
    system_cost_per_mwh = battery_cost * 1000 * 1.05 * 1.08
    initial_capex = (base_capacity * system_cost_per_mwh) / 1000000
    
    # Calculate augmentation costs
    aug_costs = {}
    for aug_year, aug_mwh in augmentations:
        aug_cost_kwh = battery_cost * ((1 + cost_decline) ** (aug_year - 1))
        aug_system_cost = aug_cost_kwh * 1000 * 1.05 * 1.08
        premium = 1.12 if aug_year < 10 else 1.15
        aug_costs[aug_year] = (aug_mwh * aug_system_cost * premium) / 1000000
    
    # Build 20-year cash flow (Sheet 3 structure)
    cf_data = []
    
    for year in range(0, project_life + 1):
        soh_pct = soh_curve.get(year, min(soh_curve.values())) / 100
        base_avail = base_capacity * soh_pct
        
        # Calculate augmented capacity
        aug_avail = 0
        for aug_year, aug_mwh in augmentations:
            if year >= aug_year:
                years_since = year - aug_year
                aug_soh = soh_curve.get(years_since, min(soh_curve.values())) / 100
                aug_avail += aug_mwh * aug_soh
        
        total_capacity = base_avail + aug_avail
        
        row = {'year': year, 'soh': soh_pct * 100, 'base_capacity': base_avail,
               'aug_capacity': aug_avail, 'total_capacity': total_capacity}
        
        if year == 0:
            row.update({'capacity_payment': 0, 'capacity_revenue': 0, 'energy_revenue': 0,
                       'total_revenue': 0, 'fixed_opex': 0, 'variable_opex': 0,
                       'total_opex': 0, 'initial_capex': -initial_capex, 'aug_capex': 0,
                       'net_cf': -initial_capex})
        else:
            if year <= 5: cp = 50000
            elif year <= 10: cp = 50000
            elif year <= 15: cp = 45000
            else: cp = 40000
            
            capacity_revenue = (power_mw * cp) / 1000000
            energy_revenue = (total_capacity * cycles_year * energy_margin) / 1000000
            total_revenue = capacity_revenue + energy_revenue
            
            fixed_opex = (850000 * (1 + inflation) ** (year - 1)) / 1000000
            
            if year <= 5: var_opex_base = 100000
            elif year <= 10: var_opex_base = 240000
            elif year <= 15: var_opex_base = 400000
            else: var_opex_base = 550000
            
            variable_opex = (var_opex_base * (1 + inflation) ** (year - 1)) / 1000000
            total_opex = fixed_opex + variable_opex
            
            aug_capex = -aug_costs.get(year, 0)
            net_cf = total_revenue - total_opex + aug_capex
            
            row.update({'capacity_payment': cp, 'capacity_revenue': capacity_revenue,
                       'energy_revenue': energy_revenue, 'total_revenue': total_revenue,
                       'fixed_opex': fixed_opex, 'variable_opex': variable_opex,
                       'total_opex': total_opex, 'initial_capex': 0, 'aug_capex': aug_capex,
                       'net_cf': net_cf})
        
        cf_data.append(row)
    
    df = pd.DataFrame(cf_data)
    df['cumulative_cf'] = df['net_cf'].cumsum()
    df['discount_factor'] = 1 / (1 + discount_rate) ** df['year']
    df['pv_cf'] = df['net_cf'] * df['discount_factor']
    df['cumulative_pv'] = df['pv_cf'].cumsum()
    
    total_capex = initial_capex
    total_revenue = df[df['year'] > 0]['total_revenue'].sum()
    total_opex = df[df['year'] > 0]['total_opex'].sum()
    npv = df['pv_cf'].sum()
    total_energy = df[df['year'] > 0]['total_capacity'].sum() * cycles_year
    lcos = ((initial_capex + sum(aug_costs.values()) + total_opex) * 1000000) / (total_energy * 1000)
    
    return {
        'name': 'Staged',
        'cashflow': df,
        'capex': initial_capex,
        'npv': npv,
        'lcos': lcos,
        'total_revenue': total_revenue,
        'total_opex': total_opex,
        'avg_capacity': df[df['year'] > 0]['total_capacity'].mean(),
        'total_energy': total_energy
    }

# ============================================================================
# REPORT GENERATION - Professional DOCX with all details
# ============================================================================

def generate_report_docx(approach, project_params, soh_curve, augmentations, results):
    """Generate comprehensive DOCX report matching Excel structure"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('BESS LCOS Analysis - Detailed Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Approach: {results['name']} | Generated: {datetime.now().strftime('%B %d, %Y')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ========== INPUT PARAMETERS ==========
    doc.add_heading('1. Input Parameters & Assumptions', level=1)
    
    doc.add_heading('Project Configuration', level=2)
    proj_table = doc.add_table(rows=6, cols=2)
    proj_table.style = 'Light Grid Accent 1'
    proj_table.rows[0].cells[0].text = 'Parameter'
    proj_table.rows[0].cells[1].text = 'Value'
    proj_table.rows[1].cells[0].text = 'Power Capacity'
    proj_table.rows[1].cells[1].text = f"{project_params['power_mw']} MW"
    proj_table.rows[2].cells[0].text = 'Energy Capacity'
    proj_table.rows[2].cells[1].text = f"{project_params['energy_mwh']} MWh"
    proj_table.rows[3].cells[0].text = 'Project Life'
    proj_table.rows[3].cells[1].text = f"{project_params['project_life']} years"
    proj_table.rows[4].cells[0].text = 'Cycles per Year'
    proj_table.rows[4].cells[1].text = f"{project_params['cycles_year']}"
    proj_table.rows[5].cells[0].text = 'Chemistry'
    proj_table.rows[5].cells[1].text = 'LiFePO4'
    
    doc.add_heading('Financial Assumptions', level=2)
    fin_table = doc.add_table(rows=7, cols=2)
    fin_table.style = 'Light Grid Accent 1'
    fin_table.rows[0].cells[0].text = 'Parameter'
    fin_table.rows[0].cells[1].text = 'Value'
    fin_table.rows[1].cells[0].text = 'Battery Cost (2025)'
    fin_table.rows[1].cells[1].text = f"${project_params['battery_cost']}/kWh"
    fin_table.rows[2].cells[0].text = 'Cost Decline Rate'
    fin_table.rows[2].cells[1].text = f"{project_params['cost_decline']*100:.2f}%/year"
    fin_table.rows[3].cells[0].text = 'Inflation Rate'
    fin_table.rows[3].cells[1].text = f"{project_params['inflation']*100:.2f}%/year"
    fin_table.rows[4].cells[0].text = 'Discount Rate'
    fin_table.rows[4].cells[1].text = f"{project_params['discount_rate']*100:.2f}%"
    fin_table.rows[5].cells[0].text = 'Energy Margin'
    fin_table.rows[5].cells[1].text = f"${project_params['energy_margin']}/MWh"
    fin_table.rows[6].cells[0].text = 'SOH Floor'
    fin_table.rows[6].cells[1].text = f"{min(soh_curve.values()):.0f}%"
    
    if augmentations:
        doc.add_heading('Augmentation Strategy', level=2)
        aug_table = doc.add_table(rows=len(augmentations)+1, cols=2)
        aug_table.style = 'Light Grid Accent 1'
        aug_table.rows[0].cells[0].text = 'Year'
        aug_table.rows[0].cells[1].text = 'Capacity (MWh)'
        for i, (year, mwh) in enumerate(augmentations, 1):
            aug_table.rows[i].cells[0].text = f"Year {year}"
            aug_table.rows[i].cells[1].text = f"{mwh}"
    
    # ========== SOH CURVE ==========
    doc.add_page_break()
    doc.add_heading('2. Battery Degradation (SOH Curve)', level=1)
    soh_table = doc.add_table(rows=len(soh_curve)+1, cols=2)
    soh_table.style = 'Light Grid Accent 1'
    soh_table.rows[0].cells[0].text = 'Year'
    soh_table.rows[0].cells[1].text = 'SOH (%)'
    for i, (year, soh) in enumerate(sorted(soh_curve.items()), 1):
        soh_table.rows[i].cells[0].text = f"Year {year}"
        soh_table.rows[i].cells[1].text = f"{soh:.1f}%"
    
    # ========== FINANCIAL RESULTS ==========
    doc.add_page_break()
    doc.add_heading('3. Financial Results Summary', level=1)
    
    res_table = doc.add_table(rows=9, cols=2)
    res_table.style = 'Light Grid Accent 1'
    res_table.rows[0].cells[0].text = 'Metric'
    res_table.rows[0].cells[1].text = 'Value'
    res_table.rows[1].cells[0].text = 'Initial CAPEX ($M)'
    res_table.rows[1].cells[1].text = f"${results['capex']:.2f}"
    res_table.rows[2].cells[0].text = 'Project NPV ($M)'
    res_table.rows[2].cells[1].text = f"${results['npv']:.2f}"
    res_table.rows[3].cells[0].text = 'LCOS ($/MWh)'
    res_table.rows[3].cells[1].text = f"${results['lcos']:.2f}"
    res_table.rows[4].cells[0].text = 'Total Revenue (20yr)'
    res_table.rows[4].cells[1].text = f"${results['total_revenue']:.2f}M"
    res_table.rows[5].cells[0].text = 'Total OPEX (20yr)'
    res_table.rows[5].cells[1].text = f"${results['total_opex']:.2f}M"
    res_table.rows[6].cells[0].text = 'Average Capacity'
    res_table.rows[6].cells[1].text = f"{results['avg_capacity']:.1f} MWh"
    res_table.rows[7].cells[0].text = 'Total Energy'
    res_table.rows[7].cells[1].text = f"{results['total_energy']:.0f} MWh"
    res_table.rows[8].cells[0].text = 'Approach'
    res_table.rows[8].cells[1].text = results['name']
    
    # ========== YEAR-BY-YEAR CASH FLOW ==========
    doc.add_page_break()
    doc.add_heading('4. Detailed Year-by-Year Cash Flow', level=1)
    
    df = results['cashflow']
    cf_table = doc.add_table(rows=min(len(df)+1, 22), cols=8)
    cf_table.style = 'Light Grid Accent 1'
    
    headers = ['Year', 'SOH %', 'Capacity MWh', 'Revenue $M', 'OPEX $M', 'CAPEX $M', 'Net CF $M', 'NPV $M']
    for col, header in enumerate(headers):
        cf_table.rows[0].cells[col].text = header
    
    for row_idx in range(1, min(len(df)+1, 22)):
        row_data = df.iloc[row_idx-1]
        cf_table.rows[row_idx].cells[0].text = f"{int(row_data['year'])}"
        cf_table.rows[row_idx].cells[1].text = f"{row_data['soh']:.1f}"
        if 'total_capacity' in df.columns:
            cf_table.rows[row_idx].cells[2].text = f"{row_data['total_capacity']:.1f}"
        else:
            cf_table.rows[row_idx].cells[2].text = f"{row_data['available_capacity']:.1f}"
        cf_table.rows[row_idx].cells[3].text = f"{row_data['total_revenue']:.2f}"
        cf_table.rows[row_idx].cells[4].text = f"{row_data['total_opex']:.2f}"
        cf_table.rows[row_idx].cells[5].text = f"{row_data['initial_capex'] + row_data['aug_capex']:.2f}"
        cf_table.rows[row_idx].cells[6].text = f"{row_data['net_cf']:.2f}"
        cf_table.rows[row_idx].cells[7].text = f"{row_data['pv_cf']:.2f}"
    
    # ========== CHARTS REFERENCE ==========
    doc.add_page_break()
    doc.add_heading('5. Charts & Visualization', level=1)
    doc.add_paragraph('Key analysis charts displayed in dashboard:')
    doc.add_paragraph('• Cumulative Cash Flow - Shows break-even analysis')
    doc.add_paragraph('• Revenue vs OPEX - Annual operational analysis')
    doc.add_paragraph('• Battery Degradation - SOH curve visualization')
    doc.add_paragraph('• Available Capacity - Capacity profile over time')
    
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

# ============================================================================
# MAIN APP
# ============================================================================

def main():
    st.markdown("<div class='header'><h1>⚡ BESS LCOS Analysis Dashboard</h1><p>Based on Professional Excel Model</p></div>", unsafe_allow_html=True)
    
    with st.sidebar:
        page = st.radio("Navigation", ["Analysis", "Reports"])
    
    if page == "Analysis":
        show_analysis()
    else:
        show_reports()

def show_analysis():
    with st.form("analysis_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Project Setup")
            power = st.number_input("Power (MW)", 50, 500, 100)
            energy = st.number_input("Energy (MWh)", 100, 2000, 400)
            life = st.number_input("Project Life (years)", 10, 40, 20)
            cycles = st.number_input("Cycles/Year", 100, 1000, 365)
        
        with col2:
            st.subheader("Financial Assumptions")
            battery = st.number_input("Battery Cost ($/kWh)", 50, 500, 241)
            decline = st.slider("Cost Decline (%/yr)", -10, 5, -4) / 100
            inflation = st.slider("Inflation (%/yr)", 0, 5, 2) / 100
            discount = st.slider("Discount (%)", 3, 15, 7) / 100
            margin = st.number_input("Energy Margin ($/MWh)", 10, 300, 50)
        
        st.markdown("---")
        st.subheader("SOH Curve Configuration")
        col1, col2, col3, col4, col5 = st.columns(5)
        with col1: soh0 = st.number_input("Y0", 100, 100, 100)
        with col2: soh5 = st.number_input("Y5", 50, 100, 88)
        with col3: soh10 = st.number_input("Y10", 50, 100, 75)
        with col4: soh15 = st.number_input("Y15", 50, 100, 63)
        with col5: soh20 = st.number_input("Y20", 50, 100, 60)
        
        soh_curve = {0: soh0, 5: soh5, 10: soh10, 15: soh15, 20: soh20}
        
        st.markdown("---")
        st.subheader("Augmentation Events")
        num_aug = st.number_input("Number of Augmentations", 0, 5, 2)
        
        augmentations = []
        if num_aug > 0:
            aug_cols = st.columns(num_aug)
            for i in range(num_aug):
                with aug_cols[i]:
                    st.markdown(f"**Aug {i+1}**")
                    yr = st.number_input("Year", 1, life-1, min(7+i*7, life-1), key=f"auy{i}")
                    mwh = st.number_input("MWh", 10, 500, 60-i*10, key=f"aum{i}")
                    augmentations.append((yr, mwh))
        
        st.markdown("---")
        st.subheader("Select Approach")
        approach = st.selectbox("Choose:", ["Initial Build (Overbuild)", "Augmentation (Staged)", "Both (Comparison)"])
        
        st.markdown("---")
        submit = st.form_submit_button("🚀 RUN ANALYSIS", type="primary", use_container_width=True)
    
    if submit:
        params = {
            'project_life': int(life), 'power_mw': power, 'energy_mwh': energy,
            'cycles_year': cycles, 'battery_cost': battery, 'cost_decline': decline,
            'inflation': inflation, 'discount_rate': discount, 'energy_margin': margin
        }
        
        try:
            if "Initial" in approach or "Both" in approach:
                st.session_state.results['overbuild'] = calculate_overbuild(params, soh_curve)
            
            if "Augmentation" in approach or "Both" in approach:
                st.session_state.results['staged'] = calculate_staged(params, soh_curve, augmentations)
            
            st.session_state.params = params
            st.session_state.soh = soh_curve
            st.session_state.augs = augmentations
            st.session_state.approach = approach
            
            st.success("✅ Analysis Complete!")
            st.rerun()
        
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
    
    # Display results
    if st.session_state.results:
        st.markdown("---")
        
        if 'overbuild' in st.session_state.results:
            ob = st.session_state.results['overbuild']
            st.markdown("<div class='header'><h2>Initial Build (Overbuild)</h2></div>", unsafe_allow_html=True)
            
            col1, col2, col3, col4 = st.columns(4)
            with col1: st.metric("CAPEX ($M)", f"${ob['capex']:.2f}")
            with col2: st.metric("NPV ($M)", f"${ob['npv']:.2f}")
            with col3: st.metric("LCOS ($/MWh)", f"${ob['lcos']:.2f}")
            with col4: st.metric("Avg Cap (MWh)", f"{ob['avg_capacity']:.1f}")
            
            tab1, tab2, tab3, tab4 = st.tabs(["📊 Cash Flow", "💹 Charts", "🔋 Capacity", "📋 Data"])
            
            with tab1:
                c1, c2 = st.columns(2)
                with c1:
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=ob['cashflow']['year'], y=ob['cashflow']['cumulative_cf'],
                                           mode='lines+markers', line=dict(color='#1976d2', width=3)))
                    fig.add_hline(y=0, line_dash='dash', line_color='red')
                    fig.update_layout(title='Cumulative Cash Flow', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='ob_cf')
                with c2:
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=ob['cashflow']['year'], y=ob['cashflow']['cumulative_pv'],
                                           mode='lines+markers', line=dict(color='#0055aa', width=3)))
                    fig.add_hline(y=0, line_dash='dash', line_color='red')
                    fig.update_layout(title='Cumulative NPV', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='ob_npv')
            
            with tab2:
                c1, c2 = st.columns(2)
                with c1:
                    df_ops = ob['cashflow'][ob['cashflow']['year'] > 0]
                    fig = go.Figure()
                    fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['total_revenue'], name='Revenue', marker_color='#4caf50'))
                    fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['total_opex'], name='OPEX', marker_color='#f44336'))
                    fig.update_layout(title='Revenue vs OPEX', barmode='group', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='ob_rev')
                with c2:
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=ob['cashflow']['year'], y=ob['cashflow']['soh'],
                                           mode='lines+markers', line=dict(color='#ff9800', width=3)))
                    fig.update_layout(title='Battery Degradation', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='ob_soh')
            
            with tab3:
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=ob['cashflow']['year'], y=ob['cashflow']['available_capacity'],
                                       mode='lines+markers', line=dict(color='#2196f3', width=3), fill='tozeroy'))
                fig.update_layout(title='Available Capacity', height=500)
                st.plotly_chart(fig, use_container_width=True, key='ob_cap')
            
            with tab4:
                st.dataframe(ob['cashflow'].round(2), use_container_width=True, height=400)
        
        if 'staged' in st.session_state.results:
            st.markdown("---")
            st_res = st.session_state.results['staged']
            st.markdown("<div class='header'><h2>Augmentation (Staged)</h2></div>", unsafe_allow_html=True)
            
            col1, col2, col3, col4 = st.columns(4)
            with col1: st.metric("CAPEX ($M)", f"${st_res['capex']:.2f}")
            with col2: st.metric("NPV ($M)", f"${st_res['npv']:.2f}")
            with col3: st.metric("LCOS ($/MWh)", f"${st_res['lcos']:.2f}")
            with col4: st.metric("Avg Cap (MWh)", f"{st_res['avg_capacity']:.1f}")
            
            tab1, tab2, tab3, tab4 = st.tabs(["📊 Cash Flow", "💹 Charts", "🔋 Capacity", "📋 Data"])
            
            with tab1:
                c1, c2 = st.columns(2)
                with c1:
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=st_res['cashflow']['year'], y=st_res['cashflow']['cumulative_cf'],
                                           mode='lines+markers', line=dict(color='#1976d2', width=3)))
                    fig.add_hline(y=0, line_dash='dash', line_color='red')
                    fig.update_layout(title='Cumulative Cash Flow', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='st_cf')
                with c2:
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=st_res['cashflow']['year'], y=st_res['cashflow']['cumulative_pv'],
                                           mode='lines+markers', line=dict(color='#0055aa', width=3)))
                    fig.add_hline(y=0, line_dash='dash', line_color='red')
                    fig.update_layout(title='Cumulative NPV', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='st_npv')
            
            with tab2:
                c1, c2 = st.columns(2)
                with c1:
                    df_ops = st_res['cashflow'][st_res['cashflow']['year'] > 0]
                    fig = go.Figure()
                    fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['total_revenue'], name='Revenue', marker_color='#4caf50'))
                    fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['total_opex'], name='OPEX', marker_color='#f44336'))
                    fig.update_layout(title='Revenue vs OPEX', barmode='group', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='st_rev')
                with c2:
                    fig = go.Figure()
                    fig.add_trace(go.Scatter(x=st_res['cashflow']['year'], y=st_res['cashflow']['soh'],
                                           mode='lines+markers', line=dict(color='#ff9800', width=3)))
                    fig.update_layout(title='Battery Degradation', height=400)
                    st.plotly_chart(fig, use_container_width=True, key='st_soh')
            
            with tab3:
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=st_res['cashflow']['year'], y=st_res['cashflow']['total_capacity'],
                                       mode='lines+markers', line=dict(color='#2196f3', width=3), fill='tozeroy'))
                fig.update_layout(title='Available Capacity', height=500)
                st.plotly_chart(fig, use_container_width=True, key='st_cap')
            
            with tab4:
                st.dataframe(st_res['cashflow'].round(2), use_container_width=True, height=400)
        
        if 'overbuild' in st.session_state.results and 'staged' in st.session_state.results:
            st.markdown("---")
            st.markdown("<div class='header'><h2>⚖️ Comparative Analysis</h2></div>", unsafe_allow_html=True)
            
            ob = st.session_state.results['overbuild']
            st_res = st.session_state.results['staged']
            
            npv_adv = st_res['npv'] - ob['npv']
            capex_adv = ob['capex'] - st_res['capex']
            lcos_adv = ob['lcos'] - st_res['lcos']
            
            col1, col2, col3, col4 = st.columns(4)
            with col1: st.metric("NPV Advantage ($M)", f"${npv_adv:.2f}", f"{(npv_adv/abs(ob['npv'])*100):.1f}%")
            with col2: st.metric("CAPEX Savings ($M)", f"${capex_adv:.2f}", f"{(capex_adv/ob['capex']*100):.1f}%")
            with col3: st.metric("LCOS Advantage ($/MWh)", f"${lcos_adv:.2f}", f"{(lcos_adv/ob['lcos']*100):.1f}%")
            with col4: st.metric("Winner", "Staged", "✓")

def show_reports():
    if not st.session_state.results:
        st.warning("⚠️ Run analysis first in Analysis tab")
        return
    
    st.markdown("<div class='header'><h2>📋 Download Detailed Reports</h2></div>", unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    if 'overbuild' in st.session_state.results:
        with col1:
            st.markdown("### Overbuild Report")
            ob = st.session_state.results['overbuild']
            docx = generate_report_docx("Overbuild", st.session_state.params, st.session_state.soh, [], ob)
            st.download_button("📥 DOCX Report", docx, f"BESS_Overbuild_{datetime.now().strftime('%Y%m%d')}.docx",
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="ob_docx")
            
            csv = ob['cashflow'].to_csv(index=False)
            st.download_button("📥 CSV Report", csv, f"BESS_Overbuild_{datetime.now().strftime('%Y%m%d')}.csv",
                             "text/csv", key="ob_csv")
    
    if 'staged' in st.session_state.results:
        with col2:
            st.markdown("### Staged Report")
            st_res = st.session_state.results['staged']
            docx = generate_report_docx("Staged", st.session_state.params, st.session_state.soh,
                                       st.session_state.augs, st_res)
            st.download_button("📥 DOCX Report", docx, f"BESS_Staged_{datetime.now().strftime('%Y%m%d')}.docx",
                             "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="st_docx")
            
            csv = st_res['cashflow'].to_csv(index=False)
            st.download_button("📥 CSV Report", csv, f"BESS_Staged_{datetime.now().strftime('%Y%m%d')}.csv",
                             "text/csv", key="st_csv")

if __name__ == "__main__":
    main()
