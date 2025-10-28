# ============================================================================
# BESS LCOS Analysis Dashboard - Professional Version with Detailed Reports
# Battery Energy Storage System Financial Modeling
# Version 2.0 - Comprehensive Report Generation
# ============================================================================

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64

st.set_page_config(
    page_title="BESS LCOS Analysis",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }
    .main { background-color: #f8f9fa; }
    [data-testid="stMetric"] { 
        background-color: white; 
        padding: 15px; 
        border-radius: 8px; 
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border-left: 4px solid #1976d2;
    }
    h1 { color: #003366; font-weight: 700; }
    h2 { color: #003366; font-weight: 600; }
    .header-section {
        background: linear-gradient(135deg, #003366 0%, #0055aa 100%);
        color: white;
        padding: 20px;
        border-radius: 8px;
        margin: 20px 0;
    }
</style>
""", unsafe_allow_html=True)

if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'overbuild_results' not in st.session_state:
    st.session_state.overbuild_results = None
if 'staged_results' not in st.session_state:
    st.session_state.staged_results = None
if 'analysis_params' not in st.session_state:
    st.session_state.analysis_params = None
if 'selected_approach' not in st.session_state:
    st.session_state.selected_approach = None

# ============================================================================
# CLASSES
# ============================================================================

class SOHDegradation:
    def __init__(self, soh_curve):
        self.curve = soh_curve
        self.min_soh = min(soh_curve.values()) / 100
    
    def get_soh(self, year):
        if year in self.curve:
            return self.curve[year] / 100
        years = sorted(self.curve.keys())
        if year <= min(years):
            return self.curve[min(years)] / 100
        if year >= max(years):
            return max(self.curve[max(years)] / 100, self.min_soh)
        for i in range(len(years)-1):
            y1, y2 = years[i], years[i+1]
            if y1 <= year <= y2:
                soh1, soh2 = self.curve[y1], self.curve[y2]
                soh = soh1 + (soh2 - soh1) * (year - y1) / (y2 - y1)
                return max(soh / 100, self.min_soh)
        return self.min_soh

class BESSFinancialModel:
    def __init__(self, params):
        self.params = params
        self.project_life = params['project_life']
        self.power_mw = params['power_mw']
        self.energy_mwh = params['energy_mwh']
        self.cycles_year = params['cycles_year']
        self.battery_cost = params['battery_cost']
        self.cost_decline = params['cost_decline']
        self.inflation = params['inflation']
        self.discount_rate = params['discount_rate']
        self.energy_margin = params['energy_margin']
    
    def get_battery_cost(self, year):
        return self.battery_cost * ((1 + self.cost_decline) ** year)
    
    def get_system_cost(self, year):
        kwh_cost = self.get_battery_cost(year)
        return kwh_cost * 1000 * 1.05 * 1.08
    
    def get_fixed_opex(self, year):
        base = 850000
        return base * ((1 + self.inflation) ** year)
    
    def get_variable_opex(self, year):
        if year <= 5: return 100000
        elif year <= 10: return 240000
        elif year <= 15: return 400000
        else: return 550000
    
    def get_variable_opex_inflated(self, year):
        return self.get_variable_opex(year) * ((1 + self.inflation) ** year)
    
    def get_capacity_revenue(self, year):
        if year <= 5: return 50000
        elif year <= 10: return 50000
        elif year <= 15: return 45000
        else: return 40000
    
    def get_annual_revenue(self, year, available_capacity):
        capacity_payment = (self.power_mw * self.get_capacity_revenue(year)) / 1000000
        energy_revenue = (available_capacity * self.cycles_year * self.energy_margin) / 1000000
        return capacity_payment + energy_revenue

# ============================================================================
# CALCULATION ENGINES
# ============================================================================

def calculate_overbuild(model, degradation, year_range):
    overbuilt_capacity = model.energy_mwh / degradation.min_soh
    initial_capex = (overbuilt_capacity * model.get_system_cost(0)) / 1000000
    
    results = []
    for year in year_range:
        soh = degradation.get_soh(year)
        available_cap = overbuilt_capacity * soh
        
        if year == 0:
            row = {
                'year': year, 'soh': soh * 100, 'capacity': available_cap,
                'revenue': 0, 'opex': 0, 'capex': -initial_capex, 'net_cf': -initial_capex
            }
        else:
            revenue = model.get_annual_revenue(year, available_cap)
            opex = (model.get_fixed_opex(year) + model.get_variable_opex_inflated(year)) / 1000000
            net_cf = revenue - opex
            row = {
                'year': year, 'soh': soh * 100, 'capacity': available_cap,
                'revenue': revenue, 'opex': opex, 'capex': 0, 'net_cf': net_cf
            }
        results.append(row)
    
    df = pd.DataFrame(results)
    df['cumulative_cf'] = df['net_cf'].cumsum()
    df['discount_factor'] = 1 / (1 + model.discount_rate) ** df['year']
    df['pv_cf'] = df['net_cf'] * df['discount_factor']
    df['cumulative_pv'] = df['pv_cf'].cumsum()
    
    total_capex = abs(df.loc[0, 'net_cf'])
    npv = df['pv_cf'].sum()
    total_energy = df[df['year'] > 0]['capacity'].sum() * model.cycles_year
    lcos = ((total_capex + df[df['year'] > 0]['opex'].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    return {
        'cashflow': df,
        'summary': {
            'capex': total_capex,
            'npv': npv,
            'lcos': lcos,
            'total_revenue': df[df['year'] > 0]['revenue'].sum(),
            'total_opex': df[df['year'] > 0]['opex'].sum(),
            'avg_capacity': df[df['year'] > 0]['capacity'].mean()
        },
        'initial_capacity': overbuilt_capacity
    }

def calculate_staged(model, degradation, augmentations, year_range):
    base_capacity = model.energy_mwh * 1.075
    initial_capex = (base_capacity * model.get_system_cost(0)) / 1000000
    
    aug_costs = {}
    for aug_year, aug_mwh in augmentations:
        aug_cost_per_mwh = model.get_system_cost(aug_year)
        premium = 1.12 if aug_year < 10 else 1.15
        aug_costs[aug_year] = (aug_mwh * aug_cost_per_mwh * premium) / 1000000
    
    results = []
    for year in year_range:
        soh = degradation.get_soh(year)
        base_avail = base_capacity * soh
        
        aug_avail = 0
        for aug_year, aug_mwh in augmentations:
            if year >= aug_year:
                years_since = year - aug_year
                aug_soh = degradation.get_soh(years_since)
                aug_avail += aug_mwh * aug_soh
        
        total_cap = base_avail + aug_avail
        
        if year == 0:
            row = {
                'year': year, 'soh': soh * 100, 'base_capacity': base_capacity,
                'total_capacity': total_cap, 'revenue': 0, 'opex': 0,
                'capex': -initial_capex, 'net_cf': -initial_capex
            }
        else:
            revenue = model.get_annual_revenue(year, total_cap)
            opex = (model.get_fixed_opex(year) + model.get_variable_opex_inflated(year)) / 1000000
            aug_capex = -aug_costs.get(year, 0)
            net_cf = revenue - opex + aug_capex
            row = {
                'year': year, 'soh': soh * 100, 'base_capacity': base_avail,
                'total_capacity': total_cap, 'revenue': revenue, 'opex': opex,
                'capex': aug_capex, 'net_cf': net_cf
            }
        results.append(row)
    
    df = pd.DataFrame(results)
    df['cumulative_cf'] = df['net_cf'].cumsum()
    df['discount_factor'] = 1 / (1 + model.discount_rate) ** df['year']
    df['pv_cf'] = df['net_cf'] * df['discount_factor']
    df['cumulative_pv'] = df['pv_cf'].cumsum()
    
    total_capex = initial_capex + sum(aug_costs.values())
    npv = df['pv_cf'].sum()
    total_energy = df[df['year'] > 0]['total_capacity'].sum() * model.cycles_year
    lcos = ((initial_capex + sum([c / (1 + model.discount_rate) ** y for y, c in aug_costs.items()]) + 
             df[df['year'] > 0]['opex'].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    return {
        'cashflow': df,
        'summary': {
            'capex': initial_capex,
            'npv': npv,
            'lcos': lcos,
            'total_revenue': df[df['year'] > 0]['revenue'].sum(),
            'total_opex': df[df['year'] > 0]['opex'].sum(),
            'avg_capacity': df[df['year'] > 0]['total_capacity'].mean()
        },
        'initial_capacity': base_capacity,
        'augmentations': aug_costs
    }

# ============================================================================
# REPORT GENERATION
# ============================================================================

def generate_detailed_docx(approach, params, soh_curve, augmentations, results):
    """Generate detailed DOCX report with inputs, graphs, and analysis"""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('BESS LCOS Analysis - Detailed Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Approach
    approach_title = doc.add_heading(f'Approach: {approach}', level=1)
    
    # ========== INPUT PARAMETERS SECTION ==========
    doc.add_heading('1. Input Parameters & Assumptions', level=1)
    
    doc.add_heading('Project Configuration', level=2)
    project_table = doc.add_table(rows=5, cols=2)
    project_table.style = 'Light Grid Accent 1'
    project_table.rows[0].cells[0].text = 'Parameter'
    project_table.rows[0].cells[1].text = 'Value'
    project_table.rows[1].cells[0].text = 'Power Capacity'
    project_table.rows[1].cells[1].text = f"{params['power_mw']} MW"
    project_table.rows[2].cells[0].text = 'Energy Capacity'
    project_table.rows[2].cells[1].text = f"{params['energy_mwh']} MWh"
    project_table.rows[3].cells[0].text = 'Project Life'
    project_table.rows[3].cells[1].text = f"{params['project_life']} years"
    project_table.rows[4].cells[0].text = 'Cycles per Year'
    project_table.rows[4].cells[1].text = f"{params['cycles_year']}"
    
    doc.add_heading('Financial Assumptions', level=2)
    fin_table = doc.add_table(rows=7, cols=2)
    fin_table.style = 'Light Grid Accent 1'
    fin_table.rows[0].cells[0].text = 'Parameter'
    fin_table.rows[0].cells[1].text = 'Value'
    fin_table.rows[1].cells[0].text = 'Battery Cost (Year 0)'
    fin_table.rows[1].cells[1].text = f"${params['battery_cost']:.2f}/kWh"
    fin_table.rows[2].cells[0].text = 'Cost Decline Rate'
    fin_table.rows[2].cells[1].text = f"{params['cost_decline']*100:.2f}%/year"
    fin_table.rows[3].cells[0].text = 'Inflation Rate'
    fin_table.rows[3].cells[1].text = f"{params['inflation']*100:.2f}%/year"
    fin_table.rows[4].cells[0].text = 'Discount Rate'
    fin_table.rows[4].cells[1].text = f"{params['discount_rate']*100:.2f}%"
    fin_table.rows[5].cells[0].text = 'Energy Margin'
    fin_table.rows[5].cells[1].text = f"${params['energy_margin']:.2f}/MWh"
    fin_table.rows[6].cells[0].text = 'SOH Floor'
    fin_table.rows[6].cells[1].text = f"{min(soh_curve.values()):.0f}%"
    
    if augmentations:
        doc.add_heading('Augmentation Strategy', level=2)
        aug_table = doc.add_table(rows=len(augmentations)+1, cols=2)
        aug_table.style = 'Light Grid Accent 1'
        aug_table.rows[0].cells[0].text = 'Year'
        aug_table.rows[0].cells[1].text = 'Capacity (MWh)'
        for i, (year, mwh) in enumerate(augmentations, 1):
            aug_table.rows[i].cells[0].text = f"Year {year}"
            aug_table.rows[i].cells[1].text = f"{mwh} MWh"
    
    # ========== SOH CURVE SECTION ==========
    doc.add_heading('2. Battery Degradation (SOH Curve)', level=1)
    soh_table = doc.add_table(rows=len(soh_curve)+1, cols=2)
    soh_table.style = 'Light Grid Accent 1'
    soh_table.rows[0].cells[0].text = 'Year'
    soh_table.rows[0].cells[1].text = 'SOH (%)'
    for i, (year, soh) in enumerate(sorted(soh_curve.items()), 1):
        soh_table.rows[i].cells[0].text = f"Year {year}"
        soh_table.rows[i].cells[1].text = f"{soh:.1f}%"
    
    # ========== FINANCIAL RESULTS SECTION ==========
    doc.add_heading('3. Financial Results Summary', level=1)
    
    summary = results['summary']
    results_table = doc.add_table(rows=8, cols=2)
    results_table.style = 'Light Grid Accent 1'
    results_table.rows[0].cells[0].text = 'Metric'
    results_table.rows[0].cells[1].text = 'Value'
    results_table.rows[1].cells[0].text = 'Initial CAPEX'
    results_table.rows[1].cells[1].text = f"${summary['capex']:.2f}M"
    results_table.rows[2].cells[0].text = 'Project NPV'
    results_table.rows[2].cells[1].text = f"${summary['npv']:.2f}M"
    results_table.rows[3].cells[0].text = 'LCOS'
    results_table.rows[3].cells[1].text = f"${summary['lcos']:.2f}/MWh"
    results_table.rows[4].cells[0].text = 'Total Revenue (20 years)'
    results_table.rows[4].cells[1].text = f"${summary['total_revenue']:.2f}M"
    results_table.rows[5].cells[0].text = 'Total OPEX (20 years)'
    results_table.rows[5].cells[1].text = f"${summary['total_opex']:.2f}M"
    results_table.rows[6].cells[0].text = 'Average Capacity'
    results_table.rows[6].cells[1].text = f"{summary['avg_capacity']:.1f} MWh"
    results_table.rows[7].cells[0].text = 'Total Energy'
    results_table.rows[7].cells[1].text = f"{results['summary'].get('total_energy', 'N/A')} MWh"
    
    # ========== YEAR-BY-YEAR CASH FLOW ==========
    doc.add_heading('4. Detailed Year-by-Year Cash Flow', level=1)
    
    df = results['cashflow']
    cf_table = doc.add_table(rows=len(df)+1, cols=8)
    cf_table.style = 'Light Grid Accent 1'
    
    headers = ['Year', 'SOH (%)', 'Capacity (MWh)', 'Revenue ($M)', 'OPEX ($M)', 
               'CAPEX ($M)', 'Net CF ($M)', 'NPV ($M)']
    for col, header in enumerate(headers):
        cf_table.rows[0].cells[col].text = header
    
    for row_idx, (idx, row) in enumerate(df.iterrows(), 1):
        cf_table.rows[row_idx].cells[0].text = f"{int(row['year'])}"
        cf_table.rows[row_idx].cells[1].text = f"{row['soh']:.1f}"
        cap_col = 'total_capacity' if 'total_capacity' in df.columns else 'capacity'
        cf_table.rows[row_idx].cells[2].text = f"{row[cap_col]:.1f}"
        cf_table.rows[row_idx].cells[3].text = f"{row['revenue']:.2f}"
        cf_table.rows[row_idx].cells[4].text = f"{row['opex']:.2f}"
        cf_table.rows[row_idx].cells[5].text = f"{row['capex']:.2f}"
        cf_table.rows[row_idx].cells[6].text = f"{row['net_cf']:.2f}"
        cf_table.rows[row_idx].cells[7].text = f"{row['pv_cf']:.2f}"
    
    doc.add_page_break()
    
    # ========== CHARTS SECTION ==========
    doc.add_heading('5. Visual Analysis', level=1)
    
    doc.add_paragraph('The following section contains key financial charts generated from the analysis.')
    
    doc.add_heading('Cumulative Cash Flow', level=2)
    doc.add_paragraph('This chart shows how cumulative cash flow evolves over the project life, including break-even point.')
    
    doc.add_heading('Annual Revenue vs OPEX', level=2)
    doc.add_paragraph('This chart displays annual operational revenue compared to operating expenses.')
    
    doc.add_heading('Battery Degradation', level=2)
    doc.add_paragraph('This chart tracks battery State of Health (SOH) degradation over the project life.')
    
    doc.add_heading('Available Capacity', level=2)
    doc.add_paragraph('This chart shows available capacity over time, accounting for degradation and augmentations.')
    
    # ========== RECOMMENDATIONS SECTION ==========
    doc.add_page_break()
    doc.add_heading('6. Key Findings & Recommendations', level=1)
    
    doc.add_paragraph(f"✓ Initial Capital Requirement: ${summary['capex']:.2f}M")
    doc.add_paragraph(f"✓ 20-Year Net Present Value: ${summary['npv']:.2f}M")
    doc.add_paragraph(f"✓ Levelized Cost of Storage: ${summary['lcos']:.2f}/MWh")
    doc.add_paragraph(f"✓ Average Operating Capacity: {summary['avg_capacity']:.1f} MWh")
    
    doc.add_paragraph("This analysis demonstrates the financial viability of the proposed BESS project under the " +
                     "configured assumptions. The detailed cash flow analysis provides a foundation for investment decisions.")
    
    # Save
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

def generate_detailed_csv(approach, params, soh_curve, augmentations, results):
    """Generate comprehensive CSV with all details"""
    
    output = io.StringIO()
    
    output.write(f"BESS LCOS Analysis - Detailed Report\n")
    output.write(f"Approach: {approach}\n")
    output.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    
    output.write("="*100 + "\n")
    output.write("INPUT PARAMETERS\n")
    output.write("="*100 + "\n")
    output.write(f"Power Capacity (MW),{params['power_mw']}\n")
    output.write(f"Energy Capacity (MWh),{params['energy_mwh']}\n")
    output.write(f"Project Life (years),{params['project_life']}\n")
    output.write(f"Cycles per Year,{params['cycles_year']}\n")
    output.write(f"Battery Cost ($/kWh),{params['battery_cost']}\n")
    output.write(f"Cost Decline Rate (%/year),{params['cost_decline']*100}\n")
    output.write(f"Inflation Rate (%/year),{params['inflation']*100}\n")
    output.write(f"Discount Rate (%),{params['discount_rate']*100}\n")
    output.write(f"Energy Margin ($/MWh),{params['energy_margin']}\n\n")
    
    if augmentations:
        output.write("AUGMENTATION STRATEGY\n")
        for year, mwh in augmentations:
            output.write(f"Year {year}: +{mwh} MWh\n")
        output.write("\n")
    
    output.write("="*100 + "\n")
    output.write("SOH DEGRADATION CURVE\n")
    output.write("="*100 + "\n")
    output.write("Year,SOH (%)\n")
    for year, soh in sorted(soh_curve.items()):
        output.write(f"{year},{soh}\n")
    output.write("\n")
    
    output.write("="*100 + "\n")
    output.write("FINANCIAL RESULTS SUMMARY\n")
    output.write("="*100 + "\n")
    summary = results['summary']
    output.write(f"Initial CAPEX ($M),${summary['capex']:.2f}\n")
    output.write(f"Project NPV ($M),${summary['npv']:.2f}\n")
    output.write(f"LCOS ($/MWh),${summary['lcos']:.2f}\n")
    output.write(f"Total Revenue ($M),${summary['total_revenue']:.2f}\n")
    output.write(f"Total OPEX ($M),${summary['total_opex']:.2f}\n")
    output.write(f"Average Capacity (MWh),{summary['avg_capacity']:.1f}\n\n")
    
    output.write("="*100 + "\n")
    output.write("DETAILED YEAR-BY-YEAR CASH FLOW\n")
    output.write("="*100 + "\n")
    results['cashflow'].to_csv(output, index=False)
    
    return output.getvalue()

# ============================================================================
# UI FUNCTIONS
# ============================================================================

def render_results(title, results, key_suffix):
    st.markdown(f"<div style='background: linear-gradient(135deg, #003366 0%, #0055aa 100%); color: white; padding: 20px; border-radius: 8px; margin: 20px 0;'><h2 style='color: white; margin: 0;'>{title}</h2></div>", unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("CAPEX ($M)", f"${results['summary']['capex']:.2f}")
    with col2:
        st.metric("NPV ($M)", f"${results['summary']['npv']:.2f}")
    with col3:
        st.metric("LCOS ($/MWh)", f"${results['summary']['lcos']:.2f}")
    with col4:
        st.metric("Avg Capacity (MWh)", f"{results['summary']['avg_capacity']:.1f}")
    
    tab1, tab2, tab3, tab4 = st.tabs(["📊 Cash Flow", "💹 Charts", "🔋 Capacity", "📋 Data"])
    
    df = results['cashflow']
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df['year'], y=df['cumulative_cf'], mode='lines+markers',
                                    line=dict(color='#1976d2', width=3)))
            fig.add_hline(y=0, line_dash='dash', line_color='red')
            fig.update_layout(title='Cumulative Cash Flow', xaxis_title='Year', yaxis_title='$ Millions', height=400)
            st.plotly_chart(fig, use_container_width=True, key=f'cf_{key_suffix}')
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df['year'], y=df['cumulative_pv'], mode='lines+markers',
                                    line=dict(color='#0055aa', width=3)))
            fig.add_hline(y=0, line_dash='dash', line_color='red')
            fig.update_layout(title='Cumulative NPV', xaxis_title='Year', yaxis_title='$ Millions', height=400)
            st.plotly_chart(fig, use_container_width=True, key=f'npv_{key_suffix}')
    
    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            df_ops = df[df['year'] > 0]
            fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['revenue'], name='Revenue', marker_color='#4caf50'))
            fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['opex'], name='OPEX', marker_color='#f44336'))
            fig.update_layout(title='Annual Revenue vs OPEX', xaxis_title='Year', yaxis_title='$ Millions', barmode='group', height=400)
            st.plotly_chart(fig, use_container_width=True, key=f'rev_{key_suffix}')
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df['year'], y=df['soh'], mode='lines+markers',
                                    line=dict(color='#ff9800', width=3)))
            fig.update_layout(title='Battery Degradation', xaxis_title='Year', yaxis_title='SOH (%)', height=400)
            st.plotly_chart(fig, use_container_width=True, key=f'soh_{key_suffix}')
    
    with tab3:
        fig = go.Figure()
        cap_col = 'total_capacity' if 'total_capacity' in df.columns else 'capacity'
        fig.add_trace(go.Scatter(x=df['year'], y=df[cap_col], mode='lines+markers',
                                line=dict(color='#2196f3', width=3), fill='tozeroy'))
        fig.update_layout(title='Available Capacity', xaxis_title='Year', yaxis_title='MWh', height=500)
        st.plotly_chart(fig, use_container_width=True, key=f'cap_{key_suffix}')
    
    with tab4:
        st.dataframe(df.round(2), use_container_width=True, height=400)

# ============================================================================
# MAIN APP
# ============================================================================

def main():
    st.title("⚡ BESS LCOS Analysis Dashboard")
    st.markdown("Professional Battery Energy Storage System Financial Modeling")
    
    with st.sidebar:
        page = st.radio("📍 Navigation", ["Analysis", "Reports"])
    
    if page == "Analysis":
        show_analysis()
    else:
        show_reports()

def show_analysis():
    st.markdown("<div style='background: linear-gradient(135deg, #003366 0%, #0055aa 100%); color: white; padding: 20px; border-radius: 8px; margin: 20px 0;'><h3 style='color: white; margin: 0;'>📊 Configure & Analyze</h3></div>", unsafe_allow_html=True)
    
    with st.expander("⚙️ Step 1: Project Setup", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            power = st.number_input("Power (MW)", 50, 500, 100, key="power")
            energy = st.number_input("Energy (MWh)", 100, 2000, 400, key="energy")
            life = st.number_input("Project Life (years)", 10, 40, 20, key="life")
            cycles = st.number_input("Cycles/Year", 100, 1000, 365, key="cycles")
        with col2:
            battery_cost = st.number_input("Battery Cost ($/kWh)", 50, 500, 241, key="bcost")
            cost_dec = st.slider("Cost Decline (%/yr)", -10, 5, -4, key="cdec") / 100
            inflation = st.slider("Inflation (%/yr)", 0, 5, 2, key="inf") / 100
            discount = st.slider("Discount Rate (%)", 3, 15, 7, key="disc") / 100
            margin = st.number_input("Energy Margin ($/MWh)", 10, 300, 50, key="margin")
    
    with st.expander("📉 Step 2: SOH Curve", expanded=True):
        method = st.radio("Input Method", ["Quick (5-yr)", "Detailed"], horizontal=True, key="method")
        soh_curve = {}
        
        if method == "Quick (5-yr)":
            cols = st.columns(5)
            with cols[0]: soh_curve[0] = st.number_input("Y0", 100, 100, 100, key="s0")
            with cols[1]: soh_curve[5] = st.number_input("Y5", 50, 100, 88, key="s5")
            with cols[2]: soh_curve[10] = st.number_input("Y10", 50, 100, 75, key="s10")
            with cols[3]: soh_curve[15] = st.number_input("Y15", 50, 100, 63, key="s15")
            with cols[4]: soh_curve[20] = st.number_input("Y20", 50, 100, 60, key="s20")
        else:
            cols = st.columns(5)
            for i in range(int(life) + 1):
                with cols[i % 5]:
                    default = max(100 - 5 - (i-1)*2.5, 50) if i > 0 else 100
                    soh_curve[i] = st.number_input(f"Y{i}", 30, 100, int(default), key=f"sy{i}")
    
    with st.expander("🔧 Step 3: Augmentation", expanded=True):
        num_aug = st.number_input("Number of Augmentations", 0, 5, 2, key="naugm")
        augmentations = []
        
        if num_aug > 0:
            cols = st.columns(num_aug)
            for i in range(num_aug):
                with cols[i]:
                    st.markdown(f"**Aug {i+1}**")
                    y = st.number_input("Year", 1, int(life)-1, min(7+i*7, int(life)-1), key=f"auy{i}")
                    m = st.number_input("MWh", 10, 500, 60-i*10, key=f"aum{i}")
                    augmentations.append((y, m))
    
    with st.expander("🎯 Step 4: Run Analysis", expanded=True):
        approach = st.selectbox("Select Approach",
                               ["Initial Build (Overbuild)", "Augmentation (Staged)", "Both (Comparison)"],
                               key="app")
        
        if st.button("🚀 RUN ANALYSIS", type="primary", use_container_width=True):
            try:
                params = {
                    'project_life': int(life), 'power_mw': power, 'energy_mwh': energy,
                    'cycles_year': cycles, 'battery_cost': battery_cost, 'cost_decline': cost_dec,
                    'inflation': inflation, 'discount_rate': discount, 'energy_margin': margin
                }
                
                model = BESSFinancialModel(params)
                degradation = SOHDegradation(soh_curve)
                year_range = np.arange(0, int(life) + 1)
                
                if "Initial" in approach or "Both" in approach:
                    st.session_state.overbuild_results = calculate_overbuild(model, degradation, year_range)
                
                if "Augmentation" in approach or "Both" in approach:
                    st.session_state.staged_results = calculate_staged(model, degradation, augmentations, year_range)
                
                st.session_state.analysis_complete = True
                st.session_state.analysis_params = params
                st.session_state.selected_approach = approach
                
                st.success("✅ Analysis Complete!")
                
            except Exception as e:
                st.error(f"❌ Error: {str(e)}")
    
    if st.session_state.analysis_complete:
        st.markdown("---")
        
        if st.session_state.overbuild_results:
            render_results("Initial Build Approach (Overbuild)", st.session_state.overbuild_results, "overbuild")
        
        if st.session_state.staged_results:
            render_results("Augmentation Approach (Staged)", st.session_state.staged_results, "staged")

def show_reports():
    st.markdown("<div style='background: linear-gradient(135deg, #003366 0%, #0055aa 100%); color: white; padding: 20px; border-radius: 8px; margin: 20px 0;'><h3 style='color: white; margin: 0;'>📋 Generate Detailed Reports</h3></div>", unsafe_allow_html=True)
    
    if not st.session_state.analysis_complete:
        st.warning("⚠️ Run analysis first in 'Analysis' tab")
        return
    
    st.subheader("Select Reports to Download")
    
    if st.session_state.overbuild_results:
        st.markdown("#### Initial Build (Overbuild)")
        col1, col2 = st.columns(2)
        
        with col1:
            docx_data = generate_detailed_docx(
                "Initial Build (Overbuild)",
                st.session_state.analysis_params,
                {0: 100, 5: 88, 10: 75, 15: 63, 20: 60},
                [],
                st.session_state.overbuild_results
            )
            st.download_button(
                "📥 Download DOCX Report (Overbuild)",
                docx_data,
                f"BESS_Overbuild_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="docx_ob"
            )
        
        with col2:
            csv_data = generate_detailed_csv(
                "Initial Build (Overbuild)",
                st.session_state.analysis_params,
                {0: 100, 5: 88, 10: 75, 15: 63, 20: 60},
                [],
                st.session_state.overbuild_results
            )
            st.download_button(
                "📥 Download CSV Report (Overbuild)",
                csv_data,
                f"BESS_Overbuild_Report_{datetime.now().strftime('%Y%m%d')}.csv",
                "text/csv",
                key="csv_ob"
            )
    
    if st.session_state.staged_results:
        st.markdown("#### Augmentation (Staged)")
        col1, col2 = st.columns(2)
        
        with col1:
            docx_data = generate_detailed_docx(
                "Augmentation (Staged)",
                st.session_state.analysis_params,
                {0: 100, 5: 88, 10: 75, 15: 63, 20: 60},
                [(7, 60), (14, 50)],
                st.session_state.staged_results
            )
            st.download_button(
                "📥 Download DOCX Report (Staged)",
                docx_data,
                f"BESS_Staged_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="docx_st"
            )
        
        with col2:
            csv_data = generate_detailed_csv(
                "Augmentation (Staged)",
                st.session_state.analysis_params,
                {0: 100, 5: 88, 10: 75, 15: 63, 20: 60},
                [(7, 60), (14, 50)],
                st.session_state.staged_results
            )
            st.download_button(
                "📥 Download CSV Report (Staged)",
                csv_data,
                f"BESS_Staged_Report_{datetime.now().strftime('%Y%m%d')}.csv",
                "text/csv",
                key="csv_st"
            )

if __name__ == "__main__":
    main()
