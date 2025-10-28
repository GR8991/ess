# ============================================================================
# BESS LCOS Analysis Dashboard - Professional Version
# Battery Energy Storage System Financial Modeling
# Author: Financial Analysis Team
# Version: 1.0 - Clean Production Build
# ============================================================================

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# PAGE CONFIGURATION
# ============================================================================

st.set_page_config(
    page_title="BESS LCOS Analysis",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; }
    .main { background-color: #f8f9fa; }
    [data-testid="stMetric"] { 
        background-color: white; 
        padding: 15px; 
        border-radius: 8px; 
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border-left: 4px solid #1976d2;
    }
    h1 { color: #003366; font-weight: 700; margin-bottom: 20px; }
    h2 { color: #003366; font-weight: 600; margin-top: 15px; }
    h3 { color: #1565c0; font-weight: 500; }
    .header-section {
        background: linear-gradient(135deg, #003366 0%, #0055aa 100%);
        color: white;
        padding: 20px;
        border-radius: 8px;
        margin: 20px 0;
        box-shadow: 0 4px 12px rgba(0,0,0,0.15);
    }
    .info-card {
        background-color: #e3f2fd;
        padding: 15px;
        border-left: 5px solid #1976d2;
        border-radius: 4px;
        margin: 15px 0;
    }
    .success-card {
        background-color: #e8f5e9;
        padding: 15px;
        border-left: 5px solid #4caf50;
        border-radius: 4px;
        margin: 15px 0;
    }
    .warning-card {
        background-color: #fff3e0;
        padding: 15px;
        border-left: 5px solid #ff9800;
        border-radius: 4px;
        margin: 15px 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

if 'analysis_complete' not in st.session_state:
    st.session_state.analysis_complete = False
if 'overbuild_results' not in st.session_state:
    st.session_state.overbuild_results = None
if 'staged_results' not in st.session_state:
    st.session_state.staged_results = None

# ============================================================================
# DATA MODELS
# ============================================================================

class SOHDegradation:
    """Battery State of Health Degradation Model"""
    def __init__(self, soh_curve):
        self.curve = soh_curve  # Dictionary: {year: soh_percentage}
        self.min_soh = min(soh_curve.values()) / 100
    
    def get_soh(self, year):
        """Get SOH for given year with linear interpolation"""
        if year in self.curve:
            return self.curve[year] / 100
        
        years = sorted(self.curve.keys())
        if year <= min(years):
            return self.curve[min(years)] / 100
        if year >= max(years):
            return max(self.curve[max(years)] / 100, self.min_soh)
        
        for i in range(len(years)-1):
            y1, y2 = years[i], years[i+1]
            if y1 <= year <= y2:
                soh1, soh2 = self.curve[y1], self.curve[y2]
                soh = soh1 + (soh2 - soh1) * (year - y1) / (y2 - y1)
                return max(soh / 100, self.min_soh)
        
        return self.min_soh

class BESSFinancialModel:
    """BESS Financial Modeling Engine"""
    
    def __init__(self, project_params):
        self.params = project_params
        self.project_life = project_params['project_life']
        self.power_mw = project_params['power_mw']
        self.energy_mwh = project_params['energy_mwh']
        self.cycles_year = project_params['cycles_year']
        self.battery_cost = project_params['battery_cost']
        self.cost_decline = project_params['cost_decline']
        self.inflation = project_params['inflation']
        self.discount_rate = project_params['discount_rate']
        self.energy_margin = project_params['energy_margin']
    
    def get_battery_cost(self, year):
        """Get battery cost for given year"""
        return self.battery_cost * ((1 + self.cost_decline) ** year)
    
    def get_system_cost(self, year):
        """Get system cost per MWh for given year"""
        kwh_cost = self.get_battery_cost(year)
        return kwh_cost * 1000 * 1.05 * 1.08  # 5% BoP, 8% EPC
    
    def get_fixed_opex(self, year):
        """Annual fixed OPEX"""
        base = 850000  # Personnel, land, maintenance, monitoring
        return base * ((1 + self.inflation) ** year)
    
    def get_variable_opex(self, year):
        """Age-dependent variable OPEX"""
        if year <= 5: return 100000
        elif year <= 10: return 240000
        elif year <= 15: return 400000
        else: return 550000
    
    def get_variable_opex_inflated(self, year):
        return self.get_variable_opex(year) * ((1 + self.inflation) ** year)
    
    def get_capacity_revenue(self, year):
        """Tiered capacity payment"""
        if year <= 5: return 50000
        elif year <= 10: return 50000
        elif year <= 15: return 45000
        else: return 40000
    
    def get_annual_revenue(self, year, available_capacity):
        """Calculate annual revenue"""
        capacity_payment = (self.power_mw * self.get_capacity_revenue(year)) / 1000000
        energy_revenue = (available_capacity * self.cycles_year * self.energy_margin) / 1000000
        return capacity_payment + energy_revenue

# ============================================================================
# CALCULATION ENGINES
# ============================================================================

def calculate_overbuild(model, degradation, year_range):
    """Calculate Initial Build (Overbuild) Approach"""
    
    overbuilt_capacity = model.energy_mwh / degradation.min_soh
    initial_capex = (overbuilt_capacity * model.get_system_cost(0)) / 1000000
    
    results = []
    for year in year_range:
        soh = degradation.get_soh(year)
        available_cap = overbuilt_capacity * soh
        
        if year == 0:
            row = {
                'year': year,
                'soh': soh * 100,
                'capacity': available_cap,
                'revenue': 0,
                'opex': 0,
                'capex': -initial_capex,
                'net_cf': -initial_capex
            }
        else:
            revenue = model.get_annual_revenue(year, available_cap)
            opex = (model.get_fixed_opex(year) + model.get_variable_opex_inflated(year)) / 1000000
            net_cf = revenue - opex
            row = {
                'year': year,
                'soh': soh * 100,
                'capacity': available_cap,
                'revenue': revenue,
                'opex': opex,
                'capex': 0,
                'net_cf': net_cf
            }
        results.append(row)
    
    df = pd.DataFrame(results)
    df['cumulative_cf'] = df['net_cf'].cumsum()
    df['discount_factor'] = 1 / (1 + model.discount_rate) ** df['year']
    df['pv_cf'] = df['net_cf'] * df['discount_factor']
    df['cumulative_pv'] = df['pv_cf'].cumsum()
    
    total_capex = abs(df.loc[0, 'net_cf'])
    npv = df['pv_cf'].sum()
    total_energy = df[df['year'] > 0]['capacity'].sum() * model.cycles_year
    lcos = ((total_capex + df[df['year'] > 0]['opex'].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    return {
        'cashflow': df,
        'summary': {
            'capex': total_capex,
            'npv': npv,
            'lcos': lcos,
            'total_revenue': df[df['year'] > 0]['revenue'].sum(),
            'total_opex': df[df['year'] > 0]['opex'].sum(),
            'avg_capacity': df[df['year'] > 0]['capacity'].mean()
        },
        'initial_capacity': overbuilt_capacity
    }

def calculate_staged(model, degradation, augmentations, year_range):
    """Calculate Augmentation (Staged) Approach"""
    
    base_capacity = model.energy_mwh * 1.075
    initial_capex = (base_capacity * model.get_system_cost(0)) / 1000000
    
    aug_dict = {year: mwh for year, mwh in augmentations}
    aug_costs = {}
    
    for aug_year, aug_mwh in augmentations:
        aug_cost_per_mwh = model.get_system_cost(aug_year)
        premium = 1.12 if aug_year < 10 else 1.15
        aug_costs[aug_year] = (aug_mwh * aug_cost_per_mwh * premium) / 1000000
    
    results = []
    for year in year_range:
        soh = degradation.get_soh(year)
        base_avail = base_capacity * soh
        
        aug_avail = 0
        for aug_year, aug_mwh in augmentations:
            if year >= aug_year:
                years_since = year - aug_year
                aug_soh = degradation.get_soh(years_since)
                aug_avail += aug_mwh * aug_soh
        
        total_cap = base_avail + aug_avail
        
        if year == 0:
            row = {
                'year': year,
                'soh': soh * 100,
                'base_capacity': base_capacity,
                'total_capacity': total_cap,
                'revenue': 0,
                'opex': 0,
                'capex': -initial_capex,
                'net_cf': -initial_capex
            }
        else:
            revenue = model.get_annual_revenue(year, total_cap)
            opex = (model.get_fixed_opex(year) + model.get_variable_opex_inflated(year)) / 1000000
            aug_capex = -aug_costs.get(year, 0)
            net_cf = revenue - opex + aug_capex
            row = {
                'year': year,
                'soh': soh * 100,
                'base_capacity': base_avail,
                'total_capacity': total_cap,
                'revenue': revenue,
                'opex': opex,
                'capex': aug_capex,
                'net_cf': net_cf
            }
        results.append(row)
    
    df = pd.DataFrame(results)
    df['cumulative_cf'] = df['net_cf'].cumsum()
    df['discount_factor'] = 1 / (1 + model.discount_rate) ** df['year']
    df['pv_cf'] = df['net_cf'] * df['discount_factor']
    df['cumulative_pv'] = df['pv_cf'].cumsum()
    
    total_capex = initial_capex + sum(aug_costs.values())
    npv = df['pv_cf'].sum()
    total_energy = df[df['year'] > 0]['total_capacity'].sum() * model.cycles_year
    lcos = ((initial_capex + sum([c / (1 + model.discount_rate) ** y for y, c in aug_costs.items()]) + 
             df[df['year'] > 0]['opex'].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    return {
        'cashflow': df,
        'summary': {
            'capex': initial_capex,
            'npv': npv,
            'lcos': lcos,
            'total_revenue': df[df['year'] > 0]['revenue'].sum(),
            'total_opex': df[df['year'] > 0]['opex'].sum(),
            'avg_capacity': df[df['year'] > 0]['total_capacity'].mean()
        },
        'initial_capacity': base_capacity,
        'augmentations': aug_costs
    }

# ============================================================================
# UI COMPONENTS
# ============================================================================

def render_results(title, results, key_suffix):
    """Render analysis results"""
    st.markdown(f"<div class='header-section'><h2>{title}</h2></div>", unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("CAPEX ($M)", f"${results['summary']['capex']:.2f}")
    with col2:
        st.metric("NPV ($M)", f"${results['summary']['npv']:.2f}")
    with col3:
        st.metric("LCOS ($/MWh)", f"${results['summary']['lcos']:.2f}")
    with col4:
        st.metric("Avg Capacity (MWh)", f"{results['summary']['avg_capacity']:.1f}")
    
    tab1, tab2, tab3, tab4 = st.tabs(["📊 Cash Flow", "💹 Charts", "🔋 Capacity", "📋 Data"])
    
    df = results['cashflow']
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df['year'], y=df['cumulative_cf'], mode='lines+markers',
                                    line=dict(color='#1976d2', width=3), name='Cumulative CF'))
            fig.add_hline(y=0, line_dash='dash', line_color='red')
            fig.update_layout(title='Cumulative Cash Flow', xaxis_title='Year', yaxis_title='$ Millions',
                            height=400, hovermode='x unified')
            st.plotly_chart(fig, use_container_width=True, key=f'cf_{key_suffix}')
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df['year'], y=df['cumulative_pv'], mode='lines+markers',
                                    line=dict(color='#0055aa', width=3), name='Cumulative NPV'))
            fig.add_hline(y=0, line_dash='dash', line_color='red')
            fig.update_layout(title='Cumulative NPV', xaxis_title='Year', yaxis_title='$ Millions',
                            height=400, hovermode='x unified')
            st.plotly_chart(fig, use_container_width=True, key=f'npv_{key_suffix}')
    
    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            df_ops = df[df['year'] > 0]
            fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['revenue'], name='Revenue', marker_color='#4caf50'))
            fig.add_trace(go.Bar(x=df_ops['year'], y=df_ops['opex'], name='OPEX', marker_color='#f44336'))
            fig.update_layout(title='Annual Revenue vs OPEX', xaxis_title='Year', yaxis_title='$ Millions',
                            barmode='group', height=400)
            st.plotly_chart(fig, use_container_width=True, key=f'rev_{key_suffix}')
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=df['year'], y=df['soh'], mode='lines+markers',
                                    line=dict(color='#ff9800', width=3), name='SOH'))
            fig.update_layout(title='Battery Degradation', xaxis_title='Year', yaxis_title='SOH (%)',
                            height=400)
            st.plotly_chart(fig, use_container_width=True, key=f'soh_{key_suffix}')
    
    with tab3:
        fig = go.Figure()
        cap_col = 'total_capacity' if 'total_capacity' in df.columns else 'capacity'
        fig.add_trace(go.Scatter(x=df['year'], y=df[cap_col], mode='lines+markers',
                                line=dict(color='#2196f3', width=3), fill='tozeroy', name='Capacity'))
        fig.update_layout(title='Available Capacity', xaxis_title='Year', yaxis_title='MWh', height=500)
        st.plotly_chart(fig, use_container_width=True, key=f'cap_{key_suffix}')
    
    with tab4:
        st.dataframe(df.round(2), use_container_width=True, height=400)

def render_comparison(overbuild, staged):
    """Render comparison analysis"""
    st.markdown("<div class='header-section'><h2>⚖️ Comparative Analysis</h2></div>", unsafe_allow_html=True)
    
    npv_adv = staged['summary']['npv'] - overbuild['summary']['npv']
    capex_adv = overbuild['summary']['capex'] - staged['summary']['capex']
    lcos_adv = overbuild['summary']['lcos'] - staged['summary']['lcos']
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("NPV Advantage", f"${npv_adv:.2f}M", f"{(npv_adv/abs(overbuild['summary']['npv'])*100):.1f}%")
    with col2:
        st.metric("CAPEX Savings", f"${capex_adv:.2f}M", f"{(capex_adv/overbuild['summary']['capex']*100):.1f}%")
    with col3:
        st.metric("LCOS Advantage", f"${lcos_adv:.2f}/MWh", f"{(lcos_adv/overbuild['summary']['lcos']*100):.1f}%")
    with col4:
        util_ob = (overbuild['summary']['avg_capacity'] / overbuild['initial_capacity'] * 100)
        util_st = (staged['summary']['avg_capacity'] / staged['initial_capacity'] * 100)
        st.metric("Better Utilization", f"{util_st:.1f}%", f"+{(util_st-util_ob):.1f}%")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        fig = go.Figure()
        fig.add_trace(go.Bar(x=['Overbuild', 'Staged'], 
                            y=[overbuild['summary']['npv'], staged['summary']['npv']],
                            marker_color=['#ff6b6b', '#51cf66']))
        fig.update_layout(title='NPV Comparison', yaxis_title='$ Millions', height=400)
        st.plotly_chart(fig, use_container_width=True, key='comp_npv')
    
    with col2:
        fig = go.Figure()
        fig.add_trace(go.Bar(x=['Overbuild', 'Staged'],
                            y=[overbuild['summary']['lcos'], staged['summary']['lcos']],
                            marker_color=['#ff6b6b', '#51cf66']))
        fig.update_layout(title='LCOS Comparison', yaxis_title='$/MWh', height=400)
        st.plotly_chart(fig, use_container_width=True, key='comp_lcos')
    
    with col3:
        fig = go.Figure()
        ob_cap = 'capacity' if 'capacity' in overbuild['cashflow'].columns else 'total_capacity'
        st_cap = 'total_capacity' if 'total_capacity' in staged['cashflow'].columns else 'capacity'
        fig.add_trace(go.Scatter(x=overbuild['cashflow']['year'], 
                                y=overbuild['cashflow'][ob_cap],
                                name='Overbuild', mode='lines', line=dict(color='red', width=2)))
        fig.add_trace(go.Scatter(x=staged['cashflow']['year'],
                                y=staged['cashflow'][st_cap],
                                name='Staged', mode='lines', line=dict(color='green', width=2)))
        fig.update_layout(title='Capacity Profile', xaxis_title='Year', yaxis_title='MWh', height=400)
        st.plotly_chart(fig, use_container_width=True, key='comp_cap')

# ============================================================================
# MAIN APPLICATION
# ============================================================================

def main():
    st.title("⚡ BESS LCOS Analysis Dashboard")
    st.markdown("Professional Battery Energy Storage System Financial Modeling")
    
    with st.sidebar:
        page = st.radio("📍 Navigation", ["Analysis", "Reports"])
    
    if page == "Analysis":
        show_analysis()
    else:
        show_reports()

def show_analysis():
    st.markdown("<div class='header-section'><h3>📊 Configure & Analyze</h3></div>", unsafe_allow_html=True)
    
    with st.expander("⚙️ Step 1: Project Setup", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            power = st.number_input("Power (MW)", 50, 500, 100, key="power")
            energy = st.number_input("Energy (MWh)", 100, 2000, 400, key="energy")
        with col2:
            life = st.number_input("Project Life (years)", 10, 40, 20, key="life")
            cycles = st.number_input("Cycles/Year", 100, 1000, 365, key="cycles")
        with col3:
            battery_cost = st.number_input("Battery Cost ($/kWh)", 50, 500, 241, key="bcost")
            cost_dec = st.slider("Cost Decline (%/yr)", -10, 5, -4, key="cdec") / 100
            inflation = st.slider("Inflation (%/yr)", 0, 5, 2, key="inf") / 100
            discount = st.slider("Discount Rate (%)", 3, 15, 7, key="disc") / 100
            margin = st.number_input("Energy Margin ($/MWh)", 10, 300, 50, key="margin")
    
    with st.expander("📉 Step 2: SOH Curve", expanded=True):
        method = st.radio("Input Method", ["Quick (5-yr)", "Detailed"], horizontal=True, key="method")
        soh_curve = {}
        
        if method == "Quick (5-yr)":
            cols = st.columns(5)
            with cols[0]: soh_curve[0] = st.number_input("Y0", 100, 100, 100, key="s0")
            with cols[1]: soh_curve[5] = st.number_input("Y5", 50, 100, 88, key="s5")
            with cols[2]: soh_curve[10] = st.number_input("Y10", 50, 100, 75, key="s10")
            with cols[3]: soh_curve[15] = st.number_input("Y15", 50, 100, 63, key="s15")
            with cols[4]: soh_curve[20] = st.number_input("Y20", 50, 100, 60, key="s20")
        else:
            cols = st.columns(5)
            for i in range(int(life) + 1):
                with cols[i % 5]:
                    default = max(100 - 5 - (i-1)*2.5, 50) if i > 0 else 100
                    soh_curve[i] = st.number_input(f"Y{i}", 30, 100, int(default), key=f"sy{i}")
        
        fig = go.Figure()
        years = sorted(soh_curve.keys())
        fig.add_trace(go.Scatter(x=years, y=[soh_curve[y] for y in years], mode='lines+markers',
                                line=dict(color='#1976d2', width=3), fill='tozeroy'))
        fig.update_layout(title='SOH Degradation Profile', xaxis_title='Year', yaxis_title='SOH (%)',
                         height=300)
        st.plotly_chart(fig, use_container_width=True, key='soh_preview')
    
    with st.expander("🔧 Step 3: Augmentation", expanded=True):
        num_aug = st.number_input("Number of Augmentations", 0, 5, 2, key="naugm")
        augmentations = []
        
        if num_aug > 0:
            cols = st.columns(num_aug)
            for i in range(num_aug):
                with cols[i]:
                    st.markdown(f"**Aug {i+1}**")
                    y = st.number_input("Year", 1, int(life)-1, min(7+i*7, int(life)-1), key=f"auy{i}")
                    m = st.number_input("MWh", 10, 500, 60-i*10, key=f"aum{i}")
                    augmentations.append((y, m))
            
            timeline = " → ".join([f"Y{y}:{m}MWh" for y, m in augmentations])
            st.markdown(f"<div class='success-card'>📅 Timeline: {timeline}</div>", unsafe_allow_html=True)
    
    with st.expander("🎯 Step 4: Run Analysis", expanded=True):
        approach = st.selectbox("Select Approach",
                               ["Initial Build (Overbuild)", "Augmentation (Staged)", "Both (Comparison)"],
                               key="app")
        
        if st.button("🚀 RUN ANALYSIS", type="primary", use_container_width=True):
            try:
                params = {
                    'project_life': int(life),
                    'power_mw': power,
                    'energy_mwh': energy,
                    'cycles_year': cycles,
                    'battery_cost': battery_cost,
                    'cost_decline': cost_dec,
                    'inflation': inflation,
                    'discount_rate': discount,
                    'energy_margin': margin
                }
                
                model = BESSFinancialModel(params)
                degradation = SOHDegradation(soh_curve)
                year_range = np.arange(0, int(life) + 1)
                
                if "Initial" in approach or "Both" in approach:
                    st.session_state.overbuild_results = calculate_overbuild(model, degradation, year_range)
                
                if "Augmentation" in approach or "Both" in approach:
                    st.session_state.staged_results = calculate_staged(model, degradation, augmentations, year_range)
                
                st.session_state.analysis_complete = True
                st.markdown("<div class='success-card'>✅ Analysis Complete!</div>", unsafe_allow_html=True)
                
            except Exception as e:
                st.markdown(f"<div class='warning-card'>❌ Error: {str(e)}</div>", unsafe_allow_html=True)
    
    if st.session_state.analysis_complete:
        st.markdown("---")
        
        if st.session_state.overbuild_results:
            render_results("Initial Build Approach (Overbuild)", st.session_state.overbuild_results, "overbuild")
        
        if st.session_state.staged_results:
            render_results("Augmentation Approach (Staged)", st.session_state.staged_results, "staged")
        
        if st.session_state.overbuild_results and st.session_state.staged_results:
            render_comparison(st.session_state.overbuild_results, st.session_state.staged_results)

def show_reports():
    st.markdown("<div class='header-section'><h3>📋 Generate Reports</h3></div>", unsafe_allow_html=True)
    
    if not st.session_state.analysis_complete:
        st.markdown("<div class='warning-card'>⚠️ Run analysis first with 'Both (Comparison)' selected</div>", unsafe_allow_html=True)
        return
    
    if not st.session_state.overbuild_results or not st.session_state.staged_results:
        st.markdown("<div class='warning-card'>⚠️ Both approaches required for reports</div>", unsafe_allow_html=True)
        return
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### 📊 CSV Report")
        csv_data = io.StringIO()
        csv_data.write("EXECUTIVE SUMMARY\n")
        csv_data.write("Metric,Overbuild,Staged,Difference\n")
        csv_data.write(f"CAPEX,${st.session_state.overbuild_results['summary']['capex']:.2f},${st.session_state.staged_results['summary']['capex']:.2f},${st.session_state.staged_results['summary']['capex']-st.session_state.overbuild_results['summary']['capex']:.2f}\n")
        csv_data.write(f"NPV,${st.session_state.overbuild_results['summary']['npv']:.2f},${st.session_state.staged_results['summary']['npv']:.2f},${st.session_state.staged_results['summary']['npv']-st.session_state.overbuild_results['summary']['npv']:.2f}\n")
        csv_data.write(f"LCOS,${st.session_state.overbuild_results['summary']['lcos']:.2f},${st.session_state.staged_results['summary']['lcos']:.2f},${st.session_state.staged_results['summary']['lcos']-st.session_state.overbuild_results['summary']['lcos']:.2f}\n")
        csv_data.seek(0)
        st.download_button("📥 Download CSV", csv_data.getvalue(),
                          f"BESS_Report_{datetime.now().strftime('%Y%m%d')}.csv",
                          "text/csv", key="csv")
    
    with col2:
        st.markdown("### 📄 DOCX Report")
        doc = Document()
        doc.add_heading('BESS LCOS Analysis Report', 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f"NPV Advantage: ${st.session_state.staged_results['summary']['npv']-st.session_state.overbuild_results['summary']['npv']:.2f}M")
        doc.add_paragraph(f"CAPEX Savings: ${st.session_state.overbuild_results['summary']['capex']-st.session_state.staged_results['summary']['capex']:.2f}M")
        
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        st.download_button("📥 Download DOCX", docx_bytes.getvalue(),
                          f"BESS_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                          key="docx")
    
    with col3:
        st.markdown("### 📊 Summary CSV")
        summary_data = pd.DataFrame({
            'Metric': ['CAPEX ($M)', 'NPV ($M)', 'LCOS ($/MWh)', 'Avg Capacity (MWh)'],
            'Overbuild': [f"{st.session_state.overbuild_results['summary']['capex']:.2f}",
                         f"{st.session_state.overbuild_results['summary']['npv']:.2f}",
                         f"{st.session_state.overbuild_results['summary']['lcos']:.2f}",
                         f"{st.session_state.overbuild_results['summary']['avg_capacity']:.1f}"],
            'Staged': [f"{st.session_state.staged_results['summary']['capex']:.2f}",
                      f"{st.session_state.staged_results['summary']['npv']:.2f}",
                      f"{st.session_state.staged_results['summary']['lcos']:.2f}",
                      f"{st.session_state.staged_results['summary']['avg_capacity']:.1f}"]
        })
        st.download_button("📥 Download Summary", summary_data.to_csv(index=False),
                          f"Summary_{datetime.now().strftime('%Y%m%d')}.csv",
                          "text/csv", key="sum")

if __name__ == "__main__":
    main()
