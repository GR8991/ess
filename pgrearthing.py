import streamlit as st
import math
from docx import Document
from docx.shared import Inches
import tempfile

st.title("*PGR* - Earth Touch & Step Voltage Calculator")

# User Inputs
c = st.number_input("Enter Earth Resistance (Ohm):", value=1.0)
g = st.selectbox("Select Ground Type:", ["Crushed", "Asphalt"])
b = st.number_input("Enter Layer Fault Duration (s):", value=1.0)
d = st.number_input("Enter Layer Thickness (m):", value=0.1)
df = st.number_input("Enter Df (Decrement Factor):", value=1.0)
sf = st.number_input("Enter Sf (Current Division Factor):", value=1.0)
If = st.number_input("Enter If (Symmetrical Ground Fault Current in kA):", value=1.0)
lc = st.number_input("Enter Lc (Total Earth Conductor for WTG & USS in m):", value=1.0)
ne = st.number_input("Enter Number of Electrodes:", step=1, value=1)
le = st.number_input("Enter Length of Each Electrode (m):", value=1.0)
area = st.number_input("Enter Area of Earth Mat (m²):", value=1.0)
h = st.number_input("Enter h (Depth of Buried Conductor in m):", value=0.5)
Lp = st.number_input("Enter Perimeter of Grid (Lp in m):", value=1.0)

# Ground type factor 'a'
a = 3000 if g == "Crushed" else 10000 if g == "Asphalt" else 0

if a == 0:
    st.error("Invalid Ground Type selected. Aborting calculation.")
else:
    # Calculate Cs
    Cs = 1 - (0.09 * (1 - (c / a))) / ((2 * d) + 0.09)
    f = round(Cs, 3)

    # Voltage Calculations
    V_70kg_touch = ((1000 + (1.5 * f * a)) * 0.157) / math.sqrt(b)
    V_50kg_touch = ((1000 + (1.5 * f * a)) * 0.116) / math.sqrt(b)
    V_70kg_step = ((1000 + (6 * f * a)) * 0.157) / math.sqrt(b)
    V_50kg_step = ((1000 + (6 * f * a)) * 0.116) / math.sqrt(b)

    # IG Calculation
    IG = df * sf * If

    # Grid Resistance Rg
    Lr = ne * le
    Lt = lc * Lr
    try:
        Rg = (
            c *
            (1 / Lt + 1 / math.sqrt(20 * area) * (1 + 1 / (1 + h * math.sqrt(20 / area)))) *
            1.52 *
            (2 * math.log(Lp * math.sqrt(2 / area) - 1) * math.sqrt(area) / Lp)
        )
    except:
        Rg = float('nan')
        st.warning("Error in Rg calculation (check inputs for valid values).")

    # Display Outputs
    st.subheader("Results")
    st.markdown("### Tolerable Touch Voltage")
    st.write(f"50 kg touch voltage: **{round(V_50kg_touch, 3)} V**")
    st.write(f"70 kg touch voltage: **{round(V_70kg_touch, 3)} V**")

    st.markdown("### Step Voltage")
    st.write(f"50 kg step voltage: **{round(V_50kg_step, 3)} V**")
    st.write(f"70 kg step voltage: **{round(V_70kg_step, 3)} V**")

    st.markdown("### Other Results")
    st.write(f"IG value: **{IG:.2f} kA**")
    st.write(f"Rg value: **{Rg:.2f} Ω**")

    # Word Document Generation
    if st.checkbox("Do you want to generate a Word report?"):
        doc = Document()
        doc.add_heading("PGR - Earth Touch & Step Voltage Calculator Report", 0)

        doc.add_heading("Inputs", level=1)
        inputs = [
            f"Earth Resistance: {c} Ohm",
            f"Ground Type: {g}",
            f"Layer Fault Duration: {b} s",
            f"Layer Thickness: {d} m",
            f"Df: {df}",
            f"Sf: {sf}",
            f"If: {If} kA",
            f"Lc: {lc} m",
            f"Number of Electrodes: {ne}",
            f"Length of Each Electrode: {le} m",
            f"Area of Earth Mat: {area} m²",
            f"Depth of Buried Conductor (h): {h} m",
            f"Perimeter of Grid (Lp): {Lp} m"
        ]
        for line in inputs:
            doc.add_paragraph(line)

        doc.add_heading("Results", level=1)

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Parameter'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Unit'

        result_data = [
            ("Touch Voltage (50kg)", round(V_50kg_touch, 3), "V"),
            ("Touch Voltage (70kg)", round(V_70kg_touch, 3), "V"),
            ("Step Voltage (50kg)", round(V_50kg_step, 3), "V"),
            ("Step Voltage (70kg)", round(V_70kg_step, 3), "V"),
            ("IG (Ground Fault Current)", round(IG, 3), "kA"),
            ("Rg (Grid Resistance)", round(Rg, 3), "Ohm")
        ]

        for param, value, unit in result_data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(param)
            row_cells[1].text = str(value)
            row_cells[2].text = str(unit)

        # Save and download the Word file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            with open(tmp.name, "rb") as file:
                st.success("Word report generated.")
                st.download_button(
                    label="Download Word Report",
                    data=file,
                    file_name="PGR_touch_step_voltage_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
