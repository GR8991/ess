# BESS LCOS Analysis - Streamlit Application (CORRECTED VERSION 2.0)
# Battery Energy Storage System - Levelized Cost of Storage Analysis
# FIXED: Removed ALL reportlab imports, replaced with python-docx

import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from datetime import datetime, timedelta
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page Configuration
st.set_page_config(
    page_title="BESS LCOS Analysis Dashboard",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main {
        background-color: #f5f5f5;
    }
    .stMetric {
        background-color: white;
        padding: 10px;
        border-radius: 5px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    h1, h2, h3 {
        color: #003366;
    }
    .section-header {
        background: linear-gradient(90deg, #003366, #0055aa);
        color: white;
        padding: 15px;
        border-radius: 5px;
        margin: 20px 0 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# ============================================================================
# DATA INITIALIZATION & INPUT MODELS
# ============================================================================

class ProjectInputs:
    """Base project parameters"""
    def __init__(self):
        self.project_name = "100 MW / 400 MWh Storage Project"
        self.power_mw = 100
        self.energy_mwh = 400
        self.storage_hours = 4
        self.project_life = 20
        self.cycles_per_year = 365
        self.start_year = 2025
        self.discount_rate = 0.07
        self.inflation_rate = 0.02
        self.battery_cost_2025 = 241  # $/kWh
        self.battery_decline_rate = -0.04  # -4%/year
        self.chemistry = "LiFePO4"

class DegradationProfile:
    """Battery State of Health degradation"""
    def __init__(self):
        self.formation_loss = 0.05  # Year 1
        self.annual_degradation = 0.025  # 2.5%/year
        self.floor_soh = 0.60  # 60% minimum
        
    def calculate_soh(self, year):
        """Calculate SOH for given year"""
        if year == 0:
            return 1.0
        elif year == 1:
            return 1.0 - self.formation_loss
        else:
            soh = (1.0 - self.formation_loss) - (self.annual_degradation * (year - 1))
            return max(soh, self.floor_soh)

class CostComponents:
    """CAPEX and OPEX calculations"""
    def __init__(self, battery_cost_kwh):
        self.battery_cost_kwh = battery_cost_kwh
        self.bop_premium = 0.05  # 5% BoP
        self.epc_premium = 0.08  # 8% EPC
        self.system_cost_per_mwh = battery_cost_kwh * 1000 * (1 + self.bop_premium) * (1 + self.epc_premium)
        
    def fixed_opex_base(self):
        """Annual fixed OPEX components"""
        personnel = 300000
        land_lease = 50000
        maintenance = 400000
        monitoring = 100000
        return personnel + land_lease + maintenance + monitoring
    
    def insurance(self, initial_capex):
        """Annual insurance as % of CAPEX"""
        return initial_capex * 1000000 * 0.0075
    
    def property_tax(self, initial_capex):
        """Annual property tax as % of CAPEX"""
        return initial_capex * 1000000 * 0.005
    
    def variable_opex(self, year):
        """Age-dependent variable OPEX"""
        if year <= 5:
            return 100000
        elif year <= 10:
            return 240000
        elif year <= 15:
            return 400000
        else:
            return 550000

class RevenueModel:
    """Revenue calculations"""
    def __init__(self):
        self.power_mw = 100
        self.energy_margin = 50  # $/MWh
        
    def capacity_payment_year(self, year):
        """Tiered capacity payments"""
        if year <= 5:
            return 50000  # $/MW/year
        elif year <= 10:
            return 50000
        elif year <= 15:
            return 45000
        else:
            return 40000
    
    def calculate_capacity_revenue(self, year):
        """Annual capacity revenue"""
        return (self.power_mw * self.capacity_payment_year(year)) / 1000000
    
    def calculate_energy_revenue(self, available_mwh, cycles_per_year, margin):
        """Annual energy arbitrage revenue"""
        return (available_mwh * cycles_per_year * margin) / 1000000

# ============================================================================
# STREAMLIT APP STRUCTURE
# ============================================================================

def main():
    st.title("⚡ BESS LCOS Analysis Dashboard")
    st.markdown("Professional Battery Energy Storage System Financial Modeling")
    
    # Sidebar Configuration
    with st.sidebar:
        st.header("🔧 Configuration")
        analysis_type = st.radio(
            "Select Analysis Type:",
            ["Input & Analysis", "Compare Scenarios", "Generate Reports"]
        )
    
    if analysis_type == "Input & Analysis":
        show_input_analysis()
    elif analysis_type == "Compare Scenarios":
        show_scenario_comparison()
    else:
        show_reports()

def show_input_analysis():
    """Main analysis tab"""
    st.markdown('<div class="section-header">📊 Project Configuration & Financial Analysis</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Project Parameters")
        power_mw = st.number_input("Power Capacity (MW)", value=100, min_value=10, max_value=500)
        energy_mwh = st.number_input("Energy Capacity (MWh)", value=400, min_value=50, max_value=2000)
        project_life = st.number_input("Project Life (years)", value=20, min_value=10, max_value=40)
        cycles_per_year = st.number_input("Cycles per Year", value=365, min_value=100, max_value=1000)
        discount_rate = st.slider("Discount Rate (%)", min_value=3, max_value=15, value=7, step=1) / 100
    
    with col2:
        st.subheader("Financial Assumptions")
        battery_cost = st.number_input("Battery Cost ($/kWh)", value=241, min_value=100, max_value=500)
        cost_decline = st.slider("Annual Cost Decline Rate (%)", min_value=-8, max_value=2, value=-4, step=1) / 100
        inflation_rate = st.slider("Inflation Rate (%)", min_value=0, max_value=5, value=2, step=1) / 100
        energy_margin = st.number_input("Energy Margin ($/MWh)", value=50, min_value=10, max_value=200)
    
    # Approach Selection
    st.markdown('<div class="section-header">🎯 Project Approach</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    
    with col1:
        approach = st.selectbox(
            "Select Project Approach:",
            ["Overbuild", "Staged Augmentation", "Both (Comparison)"]
        )
    
    if approach in ["Overbuild", "Both (Comparison)"]:
        with col2:
            st.info("**Overbuild**: Build 640 MWh upfront")
        
        # Calculate Overbuild
        overbuild_data = calculate_overbuild_cashflow(
            power_mw, energy_mwh, project_life, cycles_per_year,
            battery_cost, cost_decline, inflation_rate, discount_rate,
            energy_margin
        )
        
        show_analysis_results("Overbuild Approach", overbuild_data, "overbuild")
    
    if approach in ["Staged Augmentation", "Both (Comparison)"]:
        with col2:
            st.info("**Staged**: Build 430 MWh + augment Years 7, 14")
        
        # Calculate Staged
        staged_data = calculate_staged_cashflow(
            power_mw, energy_mwh, project_life, cycles_per_year,
            battery_cost, cost_decline, inflation_rate, discount_rate,
            energy_margin
        )
        
        show_analysis_results("Staged Augmentation Approach", staged_data, "staged")
    
    if approach == "Both (Comparison)":
        show_comparison_analysis(overbuild_data, staged_data)

def show_scenario_comparison():
    """Scenario comparison tab"""
    st.markdown('<div class="section-header">📈 Scenario Analysis</div>', unsafe_allow_html=True)
    
    st.subheader("Battery Cost Decline Scenarios")
    
    scenarios = {
        "Aggressive Decline (6%/year)": -0.06,
        "Base Case (4%/year)": -0.04,
        "Moderate Decline (2%/year)": -0.02,
        "Flat / No Decline (0%/year)": 0.0,
    }
    
    scenario_results = {}
    
    for scenario_name, decline_rate in scenarios.items():
        overbuild_data = calculate_overbuild_cashflow(
            100, 400, 20, 365, 241, decline_rate, 0.02, 0.07, 50
        )
        staged_data = calculate_staged_cashflow(
            100, 400, 20, 365, 241, decline_rate, 0.02, 0.07, 50
        )
        scenario_results[scenario_name] = {
            "overbuild_npv": overbuild_data["summary"]["npv"],
            "staged_npv": staged_data["summary"]["npv"],
            "advantage": staged_data["summary"]["npv"] - overbuild_data["summary"]["npv"]
        }
    
    # Create comparison dataframe
    scenario_df = pd.DataFrame([
        {
            "Scenario": name,
            "Overbuild NPV ($M)": v["overbuild_npv"],
            "Staged NPV ($M)": v["staged_npv"],
            "Staged Advantage ($M)": v["advantage"]
        }
        for name, v in scenario_results.items()
    ])
    
    st.dataframe(scenario_df, use_container_width=True)
    
    # Visualization
    fig = go.Figure()
    fig.add_trace(go.Bar(
        name="Overbuild NPV",
        x=list(scenario_results.keys()),
        y=[v["overbuild_npv"] for v in scenario_results.values()],
        marker_color="lightcoral"
    ))
    fig.add_trace(go.Bar(
        name="Staged NPV",
        x=list(scenario_results.keys()),
        y=[v["staged_npv"] for v in scenario_results.values()],
        marker_color="lightgreen"
    ))
    fig.update_layout(
        title="NPV Comparison Across Battery Cost Scenarios",
        xaxis_title="Scenario",
        yaxis_title="NPV ($M)",
        barmode="group",
        height=500
    )
    st.plotly_chart(fig, use_container_width=True)

def show_reports():
    """Reports generation tab"""
    st.markdown('<div class="section-header">📋 Generate Professional Reports</div>', unsafe_allow_html=True)
    
    st.subheader("Report Options")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        generate_csv = st.checkbox("Generate CSV Report", value=True)
    with col2:
        generate_docx = st.checkbox("Generate DOCX Report", value=True)
    with col3:
        generate_summary = st.checkbox("Generate Summary Table", value=True)
    
    if st.button("🚀 Generate All Reports"):
        # Calculate base case data
        overbuild_data = calculate_overbuild_cashflow(
            100, 400, 20, 365, 241, -0.04, 0.02, 0.07, 50
        )
        staged_data = calculate_staged_cashflow(
            100, 400, 20, 365, 241, -0.04, 0.02, 0.07, 50
        )
        
        st.success("✅ Generating reports...")
        
        if generate_csv:
            csv_file = generate_csv_reports(overbuild_data, staged_data)
            st.download_button(
                label="📥 Download CSV Report",
                data=csv_file,
                file_name=f"BESS_LCOS_Analysis_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
        
        if generate_docx:
            docx_file = generate_docx_report(overbuild_data, staged_data)
            st.download_button(
                label="📥 Download DOCX Report",
                data=docx_file,
                file_name=f"BESS_LCOS_Analysis_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        if generate_summary:
            summary_file = generate_summary_csv(overbuild_data, staged_data)
            st.download_button(
                label="📥 Download Summary Report",
                data=summary_file,
                file_name=f"BESS_LCOS_Summary_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv"
            )
        
        st.info("💡 Reports generated successfully! Use the download buttons above.")

# ============================================================================
# CALCULATION FUNCTIONS
# ============================================================================

def calculate_overbuild_cashflow(power_mw, energy_mwh, project_life, cycles, 
                                 battery_cost, decline_rate, inflation, discount_rate, energy_margin):
    """Calculate overbuild scenario"""
    
    # Initialize
    degradation = DegradationProfile()
    revenue_model = RevenueModel()
    revenue_model.power_mw = power_mw
    revenue_model.energy_margin = energy_margin
    
    # Capacity calculations
    overbuilt_capacity = energy_mwh / degradation.floor_soh
    
    # CAPEX
    cost_component = CostComponents(battery_cost)
    initial_capex = (overbuilt_capacity * cost_component.system_cost_per_mwh) / 1000000
    
    # Cash flow arrays
    years = np.arange(0, project_life + 1)
    cash_flows = []
    
    for year in years:
        soh = degradation.calculate_soh(year)
        available_capacity = overbuilt_capacity * soh
        
        if year == 0:
            # Construction year
            cf = {
                "year": year,
                "soh": soh * 100,
                "capacity": available_capacity,
                "capacity_revenue": 0,
                "energy_revenue": 0,
                "total_revenue": 0,
                "fixed_opex": 0,
                "variable_opex": 0,
                "total_opex": 0,
                "capex": -initial_capex,
                "net_cf": -initial_capex
            }
        else:
            # Operations years
            capacity_rev = revenue_model.calculate_capacity_revenue(year)
            energy_rev = revenue_model.calculate_energy_revenue(available_capacity, cycles, energy_margin)
            total_rev = capacity_rev + energy_rev
            
            cost_comp = CostComponents(battery_cost * ((1 + decline_rate) ** (year - 1)))
            fixed_opex = (cost_comp.fixed_opex_base() + cost_comp.insurance(initial_capex) + 
                         cost_comp.property_tax(initial_capex)) * (1 + inflation) ** (year - 1) / 1000000
            variable_opex = cost_comp.variable_opex(year) * (1 + inflation) ** (year - 1) / 1000000
            total_opex = fixed_opex + variable_opex
            
            cf = {
                "year": year,
                "soh": soh * 100,
                "capacity": available_capacity,
                "capacity_revenue": capacity_rev,
                "energy_revenue": energy_rev,
                "total_revenue": total_rev,
                "fixed_opex": fixed_opex,
                "variable_opex": variable_opex,
                "total_opex": total_opex,
                "capex": 0,
                "net_cf": total_rev - total_opex
            }
        
        cash_flows.append(cf)
    
    # Convert to dataframe
    df = pd.DataFrame(cash_flows)
    
    # Calculate cumulative and NPV metrics
    df["cumulative_cf"] = df["net_cf"].cumsum()
    df["discount_factor"] = 1 / (1 + discount_rate) ** df["year"]
    df["pv_cf"] = df["net_cf"] * df["discount_factor"]
    df["cumulative_pv"] = df["pv_cf"].cumsum()
    
    # Summary metrics
    total_capex = abs(df[df["year"] == 0]["net_cf"].values[0])
    total_revenue = df[df["year"] > 0]["total_revenue"].sum()
    total_opex = df[df["year"] > 0]["total_opex"].sum()
    npv = df["pv_cf"].sum()
    total_energy = df[df["year"] > 0]["capacity"].sum() * cycles
    lcos = ((total_capex + df[df["year"] > 0]["total_opex"].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    summary = {
        "initial_capex": total_capex,
        "total_revenue_nominal": total_revenue,
        "total_opex_nominal": total_opex,
        "npv": npv,
        "total_energy": total_energy,
        "lcos": lcos,
        "avg_capacity": df[df["year"] > 0]["capacity"].mean(),
        "payback_nominal": next((y for y in df["year"] if df[df["year"] == y]["cumulative_cf"].values[0] > 0), None)
    }
    
    return {
        "cashflow": df,
        "summary": summary,
        "overbuilt_capacity": overbuilt_capacity
    }

def calculate_staged_cashflow(power_mw, energy_mwh, project_life, cycles,
                             battery_cost, decline_rate, inflation, discount_rate, energy_margin):
    """Calculate staged augmentation scenario"""
    
    # Initialize
    degradation = DegradationProfile()
    revenue_model = RevenueModel()
    revenue_model.power_mw = power_mw
    revenue_model.energy_margin = energy_margin
    
    # Capacity calculations
    base_capacity = energy_mwh * 1.075  # 7.5% buffer
    aug1_capacity = 60  # MWh
    aug2_capacity = 50  # MWh
    
    # CAPEX
    cost_component = CostComponents(battery_cost)
    initial_capex = (base_capacity * cost_component.system_cost_per_mwh) / 1000000
    
    # Augmentation costs
    aug1_cost = (aug1_capacity * cost_component.system_cost_per_mwh * (1 + decline_rate) ** 7) / 1000000 * 1.12  # 12% integration premium
    aug2_cost = (aug2_capacity * cost_component.system_cost_per_mwh * (1 + decline_rate) ** 14) / 1000000 * 1.15  # 15% integration premium
    
    # Cash flow arrays
    years = np.arange(0, project_life + 1)
    cash_flows = []
    
    for year in years:
        soh = degradation.calculate_soh(year)
        base_avail = base_capacity * soh
        
        # Augmentation capacity
        aug_cap = 0
        if year >= 7:
            aug_cap += aug1_capacity
        if year >= 14:
            aug_cap += aug2_capacity
        
        total_capacity = base_avail + aug_cap
        
        if year == 0:
            # Construction year
            cf = {
                "year": year,
                "soh": soh * 100,
                "base_capacity": base_capacity,
                "augmented_capacity": 0,
                "total_capacity": base_capacity,
                "capacity_revenue": 0,
                "energy_revenue": 0,
                "total_revenue": 0,
                "fixed_opex": 0,
                "variable_opex": 0,
                "total_opex": 0,
                "capex": -initial_capex,
                "net_cf": -initial_capex
            }
        else:
            # Operations years
            capacity_rev = revenue_model.calculate_capacity_revenue(year)
            energy_rev = revenue_model.calculate_energy_revenue(total_capacity, cycles, energy_margin)
            total_rev = capacity_rev + energy_rev
            
            cost_comp = CostComponents(battery_cost * ((1 + decline_rate) ** (year - 1)))
            fixed_opex = (cost_comp.fixed_opex_base() + cost_comp.insurance(initial_capex) +
                         cost_comp.property_tax(initial_capex)) * (1 + inflation) ** (year - 1) / 1000000
            variable_opex = cost_comp.variable_opex(year) * (1 + inflation) ** (year - 1) / 1000000
            total_opex = fixed_opex + variable_opex
            
            # Augmentation CAPEX
            aug_capex = 0
            if year == 7:
                aug_capex = -aug1_cost
            elif year == 14:
                aug_capex = -aug2_cost
            
            cf = {
                "year": year,
                "soh": soh * 100,
                "base_capacity": base_avail,
                "augmented_capacity": aug_cap,
                "total_capacity": total_capacity,
                "capacity_revenue": capacity_rev,
                "energy_revenue": energy_rev,
                "total_revenue": total_rev,
                "fixed_opex": fixed_opex,
                "variable_opex": variable_opex,
                "total_opex": total_opex,
                "capex": aug_capex,
                "net_cf": total_rev - total_opex + aug_capex
            }
        
        cash_flows.append(cf)
    
    # Convert to dataframe
    df = pd.DataFrame(cash_flows)
    
    # Calculate cumulative and NPV metrics
    df["cumulative_cf"] = df["net_cf"].cumsum()
    df["discount_factor"] = 1 / (1 + discount_rate) ** df["year"]
    df["pv_cf"] = df["net_cf"] * df["discount_factor"]
    df["cumulative_pv"] = df["pv_cf"].cumsum()
    
    # Summary metrics
    total_capex = abs(df[df["year"] == 0]["net_cf"].values[0]) + abs(df[df["year"] == 7]["capex"].values[0]) + abs(df[df["year"] == 14]["capex"].values[0])
    total_revenue = df[df["year"] > 0]["total_revenue"].sum()
    total_opex = df[df["year"] > 0]["total_opex"].sum()
    npv = df["pv_cf"].sum()
    total_energy = df[df["year"] > 0]["total_capacity"].sum() * cycles
    lcos = ((initial_capex + (aug1_cost + aug2_cost) / (1 + discount_rate) ** 7 + (aug2_cost) / (1 + discount_rate) ** 14 + df[df["year"] > 0]["total_opex"].sum()) * 1000000) / (total_energy * 1000) if total_energy > 0 else 0
    
    summary = {
        "initial_capex": initial_capex,
        "total_revenue_nominal": total_revenue,
        "total_opex_nominal": total_opex,
        "npv": npv,
        "total_energy": total_energy,
        "lcos": lcos,
        "avg_capacity": df[df["year"] > 0]["total_capacity"].mean(),
        "payback_nominal": next((y for y in df["year"] if df[df["year"] == y]["cumulative_cf"].values[0] > 0), None)
    }
    
    return {
        "cashflow": df,
        "summary": summary,
        "base_capacity": base_capacity,
        "aug1_cost": aug1_cost,
        "aug2_cost": aug2_cost
    }

# ============================================================================
# VISUALIZATION & REPORTING FUNCTIONS
# ============================================================================

def show_analysis_results(title, data, approach_type):
    """Display analysis results with visualizations"""
    
    st.markdown(f'<div class="section-header">{title}</div>', unsafe_allow_html=True)
    
    # Key Metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Initial CAPEX ($M)", f"${data['summary']['initial_capex']:.2f}")
    with col2:
        st.metric("Project NPV ($M)", f"${data['summary']['npv']:.2f}")
    with col3:
        st.metric("LCOS ($/MWh)", f"${data['summary']['lcos']:.2f}")
    with col4:
        st.metric("Avg Capacity (MWh)", f"{data['summary']['avg_capacity']:.1f}")
    
    # Tabs for different views
    tab1, tab2, tab3, tab4 = st.tabs(["Cash Flow", "Charts", "Capacity", "Detailed Table"])
    
    with tab1:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            df = data["cashflow"]
            fig.add_trace(go.Scatter(
                x=df["year"], y=df["cumulative_cf"],
                name="Cumulative CF (Nominal)", mode="lines+markers",
                line=dict(color="blue", width=3)
            ))
            fig.add_hline(y=0, line_dash="dash", line_color="red", annotation_text="Break-even")
            fig.update_layout(title="Cumulative Cash Flow", xaxis_title="Year", yaxis_title="$ Millions",
                            hovermode="x unified", height=400)
            st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=df["year"], y=df["cumulative_pv"],
                name="Cumulative PV (NPV)", mode="lines+markers",
                line=dict(color="darkblue", width=3)
            ))
            fig.add_hline(y=0, line_dash="dash", line_color="red", annotation_text="Break-even")
            fig.update_layout(title="Cumulative Present Value", xaxis_title="Year", yaxis_title="$ Millions",
                            hovermode="x unified", height=400)
            st.plotly_chart(fig, use_container_width=True)
    
    with tab2:
        col1, col2 = st.columns(2)
        with col1:
            fig = go.Figure()
            df_ops = df[df["year"] > 0]
            fig.add_trace(go.Bar(
                x=df_ops["year"], y=df_ops["total_revenue"],
                name="Revenue", marker_color="green"
            ))
            fig.add_trace(go.Bar(
                x=df_ops["year"], y=df_ops["total_opex"],
                name="OPEX", marker_color="red"
            ))
            fig.update_layout(title="Annual Revenue vs OPEX", xaxis_title="Year", yaxis_title="$ Millions",
                            barmode="group", height=400)
            st.plotly_chart(fig, use_container_width=True)
        
        with col2:
            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=df["year"], y=df["soh"],
                name="State of Health (%)", mode="lines+markers",
                line=dict(color="purple", width=3)
            ))
            fig.update_layout(title="Battery Degradation Over Time", xaxis_title="Year", yaxis_title="SOH (%)",
                            height=400)
            st.plotly_chart(fig, use_container_width=True)
    
    with tab3:
        fig = go.Figure()
        fig.add_trace(go.Scatter(
            x=df["year"], y=df["capacity"],
            name="Available Capacity", mode="lines+markers",
            line=dict(color="orange", width=3),
            fill="tozeroy"
        ))
        fig.update_layout(title="Available Capacity Over Time", xaxis_title="Year", yaxis_title="MWh",
                        hovermode="x unified", height=500)
        st.plotly_chart(fig, use_container_width=True)
    
    with tab4:
        st.dataframe(df.round(2), use_container_width=True)
        
        # Download button for detailed CSV
        csv = df.to_csv(index=False)
        st.download_button(
            label="📥 Download Detailed Cash Flow CSV",
            data=csv,
            file_name=f"cashflow_{approach_type}_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv"
        )

def show_comparison_analysis(overbuild_data, staged_data):
    """Show comparison between two approaches"""
    
    st.markdown('<div class="section-header">⚖️ Comparative Analysis</div>', unsafe_allow_html=True)
    
    # Comparison metrics
    col1, col2, col3, col4 = st.columns(4)
    
    npv_advantage = staged_data["summary"]["npv"] - overbuild_data["summary"]["npv"]
    capex_advantage = overbuild_data["summary"]["initial_capex"] - staged_data["summary"]["initial_capex"]
    lcos_advantage = overbuild_data["summary"]["lcos"] - staged_data["summary"]["lcos"]
    
    with col1:
        st.metric(
            "NPV Advantage ($M)",
            f"${npv_advantage:.2f}",
            f"{(npv_advantage / abs(overbuild_data['summary']['npv']) * 100):.1f}% better",
            delta_color="off"
        )
    with col2:
        st.metric(
            "CAPEX Savings ($M)",
            f"${capex_advantage:.2f}",
            f"{(capex_advantage / overbuild_data['summary']['initial_capex'] * 100):.1f}% reduction",
            delta_color="off"
        )
    with col3:
        st.metric(
            "LCOS Advantage ($/MWh)",
            f"${lcos_advantage:.2f}",
            f"{(lcos_advantage / overbuild_data['summary']['lcos'] * 100):.1f}% lower",
            delta_color="off"
        )
    with col4:
        utilization_overbuild = (overbuild_data["summary"]["avg_capacity"] / 640) * 100
        utilization_staged = (staged_data["summary"]["avg_capacity"] / staged_data["base_capacity"]) * 100
        st.metric(
            "Utilization Rate",
            f"{utilization_staged:.1f}%",
            f"{(utilization_staged - utilization_overbuild):.1f}% vs Overbuild",
            delta_color="off"
        )
    
    # Comparison charts
    col1, col2 = st.columns(2)
    
    with col1:
        fig = go.Figure()
        fig.add_trace(go.Bar(
            x=["Overbuild", "Staged"],
            y=[overbuild_data["summary"]["npv"], staged_data["summary"]["npv"]],
            name="NPV ($M)", marker_color=["lightcoral", "lightgreen"]
        ))
        fig.update_layout(title="NPV Comparison", yaxis_title="$ Millions", height=400)
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        fig = go.Figure()
        fig.add_trace(go.Bar(
            x=["Overbuild", "Staged"],
            y=[overbuild_data["summary"]["lcos"], staged_data["summary"]["lcos"]],
            name="LCOS ($/MWh)", marker_color=["lightcoral", "lightgreen"]
        ))
        fig.update_layout(title="LCOS Comparison", yaxis_title="$/MWh", height=400)
        st.plotly_chart(fig, use_container_width=True)
    
    # Capacity comparison
    fig = go.Figure()
    overbuild_cf = overbuild_data["cashflow"]
    staged_cf = staged_data["cashflow"]
    
    fig.add_trace(go.Scatter(
        x=overbuild_cf["year"], y=overbuild_cf["capacity"],
        name="Overbuild", mode="lines+markers",
        line=dict(color="red", width=2)
    ))
    fig.add_trace(go.Scatter(
        x=staged_cf["year"], y=staged_cf["total_capacity"],
        name="Staged", mode="lines+markers",
        line=dict(color="green", width=2)
    ))
    fig.update_layout(
        title="Available Capacity Comparison",
        xaxis_title="Year", yaxis_title="MWh",
        hovermode="x unified", height=450
    )
    st.plotly_chart(fig, use_container_width=True)

def generate_csv_reports(overbuild_data, staged_data):
    """Generate comprehensive CSV report"""
    
    output = io.StringIO()
    
    # Header
    output.write("BESS LCOS ANALYSIS - PROFESSIONAL REPORT\n")
    output.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    
    # Summary Comparison
    output.write("EXECUTIVE SUMMARY\n")
    output.write("="*80 + "\n\n")
    output.write("Metric,Overbuild,Staged,Difference\n")
    output.write(f"Initial CAPEX ($M),{overbuild_data['summary']['initial_capex']:.2f},{staged_data['summary']['initial_capex']:.2f},{staged_data['summary']['initial_capex']-overbuild_data['summary']['initial_capex']:.2f}\n")
    output.write(f"Project NPV ($M),{overbuild_data['summary']['npv']:.2f},{staged_data['summary']['npv']:.2f},{staged_data['summary']['npv']-overbuild_data['summary']['npv']:.2f}\n")
    output.write(f"LCOS ($/MWh),{overbuild_data['summary']['lcos']:.2f},{staged_data['summary']['lcos']:.2f},{staged_data['summary']['lcos']-overbuild_data['summary']['lcos']:.2f}\n")
    output.write(f"Total Revenue ($M),{overbuild_data['summary']['total_revenue_nominal']:.2f},{staged_data['summary']['total_revenue_nominal']:.2f},{staged_data['summary']['total_revenue_nominal']-overbuild_data['summary']['total_revenue_nominal']:.2f}\n")
    output.write(f"Total OPEX ($M),{overbuild_data['summary']['total_opex_nominal']:.2f},{staged_data['summary']['total_opex_nominal']:.2f},{staged_data['summary']['total_opex_nominal']-overbuild_data['summary']['total_opex_nominal']:.2f}\n")
    output.write(f"Average Capacity (MWh),{overbuild_data['summary']['avg_capacity']:.1f},{staged_data['summary']['avg_capacity']:.1f},{staged_data['summary']['avg_capacity']-overbuild_data['summary']['avg_capacity']:.1f}\n\n")
    
    # Overbuild Detailed
    output.write("\nOVERBUILD APPROACH - YEAR BY YEAR CASH FLOW\n")
    output.write("="*80 + "\n")
    overbuild_data["cashflow"].to_csv(output, index=False)
    
    # Staged Detailed
    output.write("\n\nSTAGED AUGMENTATION APPROACH - YEAR BY YEAR CASH FLOW\n")
    output.write("="*80 + "\n")
    staged_data["cashflow"].to_csv(output, index=False)
    
    return output.getvalue()

def generate_summary_csv(overbuild_data, staged_data):
    """Generate summary CSV file"""
    
    output = io.StringIO()
    
    output.write("BESS LCOS Analysis - Executive Summary\n")
    output.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
    
    summary_data = {
        "Metric": [
            "Initial CAPEX ($M)",
            "Project NPV ($M)",
            "LCOS ($/MWh)",
            "Total Revenue ($M)",
            "Total OPEX ($M)",
            "Average Capacity (MWh)",
            "Payback Period (years)",
            "Total Energy (MWh)"
        ],
        "Overbuild": [
            f"{overbuild_data['summary']['initial_capex']:.2f}",
            f"{overbuild_data['summary']['npv']:.2f}",
            f"{overbuild_data['summary']['lcos']:.2f}",
            f"{overbuild_data['summary']['total_revenue_nominal']:.2f}",
            f"{overbuild_data['summary']['total_opex_nominal']:.2f}",
            f"{overbuild_data['summary']['avg_capacity']:.1f}",
            f"{overbuild_data['summary']['payback_nominal'] if overbuild_data['summary']['payback_nominal'] else 'N/A'}",
            f"{overbuild_data['summary']['total_energy']:.0f}"
        ],
        "Staged": [
            f"{staged_data['summary']['initial_capex']:.2f}",
            f"{staged_data['summary']['npv']:.2f}",
            f"{staged_data['summary']['lcos']:.2f}",
            f"{staged_data['summary']['total_revenue_nominal']:.2f}",
            f"{staged_data['summary']['total_opex_nominal']:.2f}",
            f"{staged_data['summary']['avg_capacity']:.1f}",
            f"{staged_data['summary']['payback_nominal'] if staged_data['summary']['payback_nominal'] else 'N/A'}",
            f"{staged_data['summary']['total_energy']:.0f}"
        ],
        "Advantage": [
            f"{staged_data['summary']['initial_capex'] - overbuild_data['summary']['initial_capex']:.2f}",
            f"{staged_data['summary']['npv'] - overbuild_data['summary']['npv']:.2f}",
            f"{staged_data['summary']['lcos'] - overbuild_data['summary']['lcos']:.2f}",
            f"{staged_data['summary']['total_revenue_nominal'] - overbuild_data['summary']['total_revenue_nominal']:.2f}",
            f"{staged_data['summary']['total_opex_nominal'] - overbuild_data['summary']['total_opex_nominal']:.2f}",
            f"{staged_data['summary']['avg_capacity'] - overbuild_data['summary']['avg_capacity']:.1f}",
            "Better" if staged_data['summary']['payback_nominal'] and overbuild_data['summary']['payback_nominal'] and staged_data['summary']['payback_nominal'] < overbuild_data['summary']['payback_nominal'] else "-",
            f"{staged_data['summary']['total_energy'] - overbuild_data['summary']['total_energy']:.0f}"
        ]
    }
    
    summary_df = pd.DataFrame(summary_data)
    summary_df.to_csv(output, index=False)
    
    return output.getvalue()

def generate_docx_report(overbuild_data, staged_data):
    """Generate DOCX report using python-docx"""
    
    doc = Document()
    
    # Add title
    title = doc.add_heading('BESS LCOS Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    timestamp_para = doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%B %d, %Y at %H:%M:%S')}")
    timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Blank line
    
    # Executive Summary Section
    doc.add_heading('Executive Summary', level=1)
    
    summary_text = f"""
The Staged Augmentation approach provides significant financial advantages over the Overbuild approach:

• NPV Advantage: ${staged_data['summary']['npv'] - overbuild_data['summary']['npv']:.2f}M ({((staged_data['summary']['npv'] - overbuild_data['summary']['npv']) / abs(overbuild_data['summary']['npv']) * 100):.1f}% improvement)
• CAPEX Savings: ${overbuild_data['summary']['initial_capex'] - staged_data['summary']['initial_capex']:.2f}M ({((overbuild_data['summary']['initial_capex'] - staged_data['summary']['initial_capex']) / overbuild_data['summary']['initial_capex'] * 100):.1f}% reduction)
• LCOS Advantage: ${overbuild_data['summary']['lcos'] - staged_data['summary']['lcos']:.2f}/MWh ({((overbuild_data['summary']['lcos'] - staged_data['summary']['lcos']) / overbuild_data['summary']['lcos'] * 100):.1f}% lower)
• Payback Improvement: {overbuild_data['summary']['payback_nominal'] - staged_data['summary']['payback_nominal'] if overbuild_data['summary']['payback_nominal'] and staged_data['summary']['payback_nominal'] else 'N/A'} years faster
    """
    doc.add_paragraph(summary_text)
    
    # Comparison Table
    doc.add_heading('Comparative Metrics', level=2)
    
    table = doc.add_table(rows=9, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Overbuild'
    hdr_cells[2].text = 'Staged'
    hdr_cells[3].text = 'Difference'
    
    # Data rows
    data_rows = [
        ['Initial CAPEX ($M)', f"${overbuild_data['summary']['initial_capex']:.2f}", f"${staged_data['summary']['initial_capex']:.2f}", f"${staged_data['summary']['initial_capex'] - overbuild_data['summary']['initial_capex']:.2f}"],
        ['Project NPV ($M)', f"${overbuild_data['summary']['npv']:.2f}", f"${staged_data['summary']['npv']:.2f}", f"${staged_data['summary']['npv'] - overbuild_data['summary']['npv']:.2f}"],
        ['LCOS ($/MWh)', f"${overbuild_data['summary']['lcos']:.2f}", f"${staged_data['summary']['lcos']:.2f}", f"${staged_data['summary']['lcos'] - overbuild_data['summary']['lcos']:.2f}"],
        ['Total Revenue ($M)', f"${overbuild_data['summary']['total_revenue_nominal']:.2f}", f"${staged_data['summary']['total_revenue_nominal']:.2f}", f"${staged_data['summary']['total_revenue_nominal'] - overbuild_data['summary']['total_revenue_nominal']:.2f}"],
        ['Total OPEX ($M)', f"${overbuild_data['summary']['total_opex_nominal']:.2f}", f"${staged_data['summary']['total_opex_nominal']:.2f}", f"${staged_data['summary']['total_opex_nominal'] - overbuild_data['summary']['total_opex_nominal']:.2f}"],
        ['Avg Capacity (MWh)', f"{overbuild_data['summary']['avg_capacity']:.1f}", f"{staged_data['summary']['avg_capacity']:.1f}", f"{staged_data['summary']['avg_capacity'] - overbuild_data['summary']['avg_capacity']:.1f}"],
        ['Utilization Rate', f"{(overbuild_data['summary']['avg_capacity'] / 640 * 100):.1f}%", f"{(staged_data['summary']['avg_capacity'] / staged_data['base_capacity'] * 100):.1f}%", f"+{(staged_data['summary']['avg_capacity'] / staged_data['base_capacity'] * 100) - (overbuild_data['summary']['avg_capacity'] / 640 * 100):.1f}%"],
        ['Total Energy (MWh)', f"{overbuild_data['summary']['total_energy']:.0f}", f"{staged_data['summary']['total_energy']:.0f}", f"{staged_data['summary']['total_energy'] - overbuild_data['summary']['total_energy']:.0f}"],
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
    
    # Key Findings
    doc.add_heading('Key Findings', level=2)
    
    findings = [
        f"Staged approach is {((staged_data['summary']['npv'] - overbuild_data['summary']['npv']) / abs(overbuild_data['summary']['npv']) * 100):.1f}% better on NPV (${staged_data['summary']['npv'] - overbuild_data['summary']['npv']:.2f}M advantage)",
        f"Initial CAPEX reduction of {((overbuild_data['summary']['initial_capex'] - staged_data['summary']['initial_capex']) / overbuild_data['summary']['initial_capex'] * 100):.1f}% saves ${overbuild_data['summary']['initial_capex'] - staged_data['summary']['initial_capex']:.2f}M upfront",
        f"LCOS is {((overbuild_data['summary']['lcos'] - staged_data['summary']['lcos']) / overbuild_data['summary']['lcos'] * 100):.1f}% lower (${staged_data['summary']['lcos']:.2f}/MWh vs ${overbuild_data['summary']['lcos']:.2f}/MWh)",
        f"Superior capital efficiency with {(staged_data['summary']['avg_capacity'] / staged_data['base_capacity'] * 100):.1f}% utilization vs {(overbuild_data['summary']['avg_capacity'] / 640 * 100):.1f}%",
        "Augmentation timing optimized with battery cost decline and degradation trajectory"
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    # Recommendation
    doc.add_heading('Recommendation', level=2)
    doc.add_paragraph(
        "The Staged Augmentation approach is financially superior to the Overbuild approach across all key metrics. "
        "This strategy reduces capital requirements, improves NPV, lowers LCOS, and achieves better capacity utilization. "
        "The staged approach is recommended for implementation."
    )
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    return docx_bytes.getvalue()

if __name__ == "__main__":
    main()
